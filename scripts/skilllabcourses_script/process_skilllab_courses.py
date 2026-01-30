#!/usr/bin/env python3
"""
Process Skill Lab Courses from DOCX files
- Extract MCQs from chapters and create separate MCQ docx files
- Mark chapter name as Heading1 and sections as Heading2
- Generate HTML for chapters (with text between H1 and first H2 commented out)
- Generate HTML for intro.docx
- Generate PDF for worksheet.docx and MCQ.docx
"""

import os
import sys
import re
from pathlib import Path
from typing import List, Tuple, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# Add parent directory to path to import convert_docx_to_html
scripts_dir = Path(__file__).parent.parent
sys.path.insert(0, str(scripts_dir))
# Add current directory for pdf_templates
sys.path.insert(0, str(Path(__file__).parent))
try:
    from convert_docx_to_html import convert_docx_to_html
except ImportError:
    print("Warning: Could not import convert_docx_to_html. HTML conversion may fail.")
    print(f"  Tried to import from: {scripts_dir}")
    def convert_docx_to_html(docx_path):
        return None

try:
    import pdfkit
    PDFKIT_AVAILABLE = True
    WEASYPRINT_AVAILABLE = False
except ImportError:
    try:
        from weasyprint import HTML as WeasyHTML
        PDFKIT_AVAILABLE = False
        WEASYPRINT_AVAILABLE = True
    except ImportError:
        PDFKIT_AVAILABLE = False
        WEASYPRINT_AVAILABLE = False
        print("Warning: Neither pdfkit nor weasyprint available. PDF generation will be skipped.")

# Configuration
SOURCE_DIR = Path("/home/itpc6/Public/share/content- Topteen/skill lab courses")
OUTPUT_DIR = Path("/home/itpc6/Public/django/git-repo/7nov/git/new_template-demo-topteens/topteen_1.0/skilllabcourses_html")
SCRIPT_DIR = Path("/home/itpc6/Public/django/git-repo/7nov/git/new_template-demo-topteens/topteen_1.0/scripts/skilllabcourses_script")


def extract_chapter_name_from_docx(chapter_docx_path: Path) -> Optional[str]:
    """
    Extract Heading1 (chapter name) from chapter DOCX file.
    Returns the first Heading1 text found, or None if not found.
    """
    try:
        doc = Document(str(chapter_docx_path))
        
        # First, check for explicit Heading 1 style
        for para in doc.paragraphs[:30]:  # Check first 30 paragraphs
            if para.style.name == 'Heading 1':
                text = para.text.strip()
                if text:
                    return text
        
        # If no explicit Heading1, look for bold paragraphs that look like headings
        for para in doc.paragraphs[:30]:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if it starts with "Chapter" and is bold
            if text.startswith('Chapter') and para.runs:
                # Check if first run is bold
                if para.runs[0].bold:
                    return text
            
            # Check if paragraph is bold and looks like a heading (short, starts with capital)
            if len(text) < 200 and text[0].isupper() and para.runs:
                if any(run.bold for run in para.runs):
                    # Check if it matches chapter pattern
                    if re.match(r'^Chapter\s+\d+', text, re.IGNORECASE):
                        return text
        
        return None
    except Exception as e:
        print(f"    Error extracting chapter name from {chapter_docx_path.name}: {e}")
        return None

# Global tracking for questions without correct answers
MISSING_ANSWERS_TRACKER = []


def detect_mcq_section(paragraphs: List) -> Tuple[int, int]:
    """
    Detect MCQ section in chapter document.
    Returns (start_index, end_index) of MCQ section, or (-1, -1) if not found.
    Looks for MCQ patterns and reads till end of document.
    """
    start_idx = -1
    end_idx = -1
    
    # Look for common MCQ indicators
    mcq_indicators = [
        r'^MCQ',
        r'^Multiple Choice',
        r'^MCQs',
        r'^Questions',
        r'^Exercise',
        r'^Assessment',
        r'^Quiz',
        r'^Practice Questions',
        r'^Review Questions',
        r'^Self-Assessment',
        r'^Test Your Knowledge',
    ]
    
    # Also look for question patterns like "Q1", "1.", "Question 1", etc.
    question_patterns = [
        r'^Q\d+',
        r'^\d+\.\s+[A-Z]',  # Numbered questions
        r'^Question\s+\d+',
        r'^\d+\)',  # Numbered with parenthesis
    ]
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        # Check if this looks like the start of MCQ section
        if start_idx == -1:
            # Check for explicit MCQ indicators
            for pattern in mcq_indicators:
                if re.match(pattern, text, re.IGNORECASE):
                    start_idx = i
                    break
            
            # If not found, check for question patterns (but only if we're past the middle of document)
            if start_idx == -1 and i > len(paragraphs) * 0.6:
                for pattern in question_patterns:
                    if re.match(pattern, text, re.IGNORECASE):
                        # Verify it's likely MCQs by checking next few paragraphs for options (a, b, c, d)
                        options_found = 0
                        for j in range(i + 1, min(i + 10, len(paragraphs))):
                            next_text = paragraphs[j].text.strip().lower()
                            if re.match(r'^[a-d][\.\)]', next_text) or re.match(r'^\([a-d]\)', next_text):
                                options_found += 1
                        if options_found >= 2:  # Found at least 2 option patterns
                            start_idx = i
                            break
    
    # If start found, MCQs go to end of document
    if start_idx != -1:
        end_idx = len(paragraphs)
    
    return (start_idx, end_idx)


def extract_mcqs_from_chapter(chapter_doc: Document) -> Optional[Document]:
    """
    Extract MCQs from chapter document and create a separate MCQ document.
    Returns new Document with MCQs, or None if no MCQs found.
    """
    paragraphs = list(chapter_doc.paragraphs)
    start_idx, end_idx = detect_mcq_section(paragraphs)
    
    if start_idx == -1:
        return None
    
    # Create new document for MCQs
    mcq_doc = Document()
    
    # Copy MCQ section
    for i in range(start_idx, end_idx):
        para = paragraphs[i]
        new_para = mcq_doc.add_paragraph()
        
        # Copy paragraph properties
        new_para.style = para.style
        new_para.alignment = para.alignment
        
        # Copy runs
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
    
    # Copy tables in MCQ section (if any)
    for table in chapter_doc.tables:
        # Check if table is in MCQ section by checking if it comes after start_idx
        # This is approximate - we'll include all tables for safety
        new_table = mcq_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
    
    return mcq_doc


def remove_mcqs_from_chapter(chapter_doc: Document) -> Document:
    """
    Remove MCQ section from chapter document.
    Returns modified document.
    """
    paragraphs = list(chapter_doc.paragraphs)
    start_idx, end_idx = detect_mcq_section(paragraphs)
    
    if start_idx == -1:
        return chapter_doc
    
    # Remove paragraphs in reverse order to maintain indices
    for i in range(end_idx - 1, start_idx - 1, -1):
        para = paragraphs[i]
        para._p.getparent().remove(para._p)
    
    return chapter_doc


def mark_headings_in_chapter(chapter_doc: Document, chapter_name: str):
    """
    Mark chapter name as Heading1 and sections as Heading2.
    """
    paragraphs = list(chapter_doc.paragraphs)
    
    # Get or create Heading styles
    try:
        heading1_style = chapter_doc.styles['Heading 1']
    except:
        heading1_style = chapter_doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    try:
        heading2_style = chapter_doc.styles['Heading 2']
    except:
        heading2_style = chapter_doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    # Find chapter name paragraph (usually first non-empty paragraph or one matching chapter name)
    chapter_found = False
    for para in paragraphs[:10]:  # Check first 10 paragraphs
        text = para.text.strip()
        if text:
            # Check if it matches chapter name pattern
            if (chapter_name.lower() in text.lower() or 
                text.startswith('Chapter') or 
                re.match(r'^Chapter\s+\d+', text, re.IGNORECASE)):
                para.style = heading1_style
                chapter_found = True
                break
    
    # If chapter name not found, mark first non-empty paragraph as Heading1
    if not chapter_found:
        for para in paragraphs:
            if para.text.strip():
                para.style = heading1_style
                break
    
    # Mark sections as Heading2
    # Look for patterns like "Section 1", "1.", "Introduction", "Conclusion", etc.
    section_patterns = [
        r'^Section\s+\d+',
        r'^\d+\.\s+[A-Z][a-z]',  # Numbered sections (e.g., "1. Introduction")
        r'^Introduction\s*$',
        r'^Conclusion\s*$',
        r'^Summary\s*$',
        r'^Overview\s*$',
        r'^In\s+Conclusion',  # "In Conclusion"
        r'^To\s+Conclude',  # "To Conclude"
        r'^Final\s+Thoughts',  # "Final Thoughts"
        r'^Key\s+Takeaways',  # "Key Takeaways"
    ]
    
    # First, mark regular sections
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Skip if already Heading1
        if para.style.name == 'Heading 1':
            continue
        
        # Check if this looks like a section heading
        is_section = False
        for pattern in section_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                is_section = True
                break
        
        # Also check formatting hints (bold, larger font, centered)
        if not is_section:
            # Check if paragraph is bold and short (likely a heading)
            if len(text) < 100:  # Short text
                for run in para.runs:
                    if run.bold and run.font.size and run.font.size.pt >= 14:
                        is_section = True
                        break
        
        if is_section:
            para.style = heading2_style

    # Now handle "Wrap-up" patterns - must be bold and near end of document
    # Look for variations: "Wrap-up", "Wrap-Up", "Wrap up", "Wrap-Up for Chapter X", etc.
    wrap_up_patterns = [
        r'.*Wrap[-\s]?[Uu]p.*',  # "Wrap-up", "Wrap-Up", "Wrap up"
        r'.*Wrap[-\s]?[Uu]p\s+for\s+Chapter.*',  # "Wrap-up for Chapter X"
        r'.*Wrap[-\s]?[Uu]p\s+Chapter.*',  # "Wrap-up Chapter X"
    ]
    
    total_paragraphs = len(paragraphs)
    # Check paragraphs from 60% to end of document (likely conclusion area)
    start_check_index = max(0, int(total_paragraphs * 0.6))
    
    for i in range(start_check_index, total_paragraphs):
        para = paragraphs[i]
        text = para.text.strip()
        if not text:
            continue
        
        # Skip if already Heading1 or Heading2
        if para.style.name in ['Heading 1', 'Heading 2']:
            continue
        
        # Check if text contains wrap-up pattern
        matches_wrap_up = False
        for pattern in wrap_up_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                matches_wrap_up = True
                break
        
        if matches_wrap_up:
            # Check if paragraph is bold (all runs or most runs)
            is_bold = False
            bold_runs = 0
            total_runs = len(para.runs) if para.runs else 0
            
            if total_runs > 0:
                for run in para.runs:
                    if run.bold:
                        bold_runs += 1
                # Consider bold if at least 50% of runs are bold or if text is short and bold
                if bold_runs >= (total_runs * 0.5) or (len(text) < 100 and bold_runs > 0):
                    is_bold = True
            else:
                # No runs, check paragraph-level formatting
                # If paragraph has no runs but has text, check if it's formatted as heading
                if len(text) < 150:  # Short text likely to be heading
                    is_bold = True
            
            # Only mark as Heading2 if it's bold and in the last portion of document
            if is_bold:
                para.style = heading2_style
                break  # Only mark the first wrap-up found near the end


def docx_to_pdf(docx_path: Path, pdf_path: Path, pdf_type: str = 'general', 
                 chapter_name: str = '', course_name: str = '') -> bool:
    """
    Convert DOCX to PDF using available library with professional formatting.
    
    Args:
        docx_path: Path to DOCX file
        pdf_path: Path to output PDF file
        pdf_type: Type of PDF ('mcq', 'worksheet', or 'general')
        chapter_name: Chapter name for header
        course_name: Course name for header
    """
    try:
        # Import PDF templates
        try:
            from pdf_templates import get_mcq_pdf_template, get_worksheet_pdf_template, get_worksheet_pdf_from_html, format_mcq_html_for_pdf
        except ImportError:
            get_mcq_pdf_template = None
            get_worksheet_pdf_template = None
            get_worksheet_pdf_from_html = None
            format_mcq_html_for_pdf = None
        
        # First convert docx to HTML
        html_content = convert_docx_to_html(docx_path)
        if not html_content:
            print(f"  Error: Could not convert {docx_path} to HTML")
            return False
        
        # Apply professional templates for MCQ and Worksheet
        if pdf_type == 'mcq' and get_mcq_pdf_template:
            # Format MCQ HTML first
            if format_mcq_html_for_pdf:
                formatted_content = format_mcq_html_for_pdf(html_content)
                # Create full template
                html_content = get_mcq_pdf_template(formatted_content, chapter_name, course_name)
            else:
                html_content = get_mcq_pdf_template(html_content, chapter_name, course_name)
        elif pdf_type == 'worksheet':
            # Use the new structured worksheet PDF function (similar to MCQ PDF)
            if get_worksheet_pdf_from_html:
                html_content = get_worksheet_pdf_from_html(html_content, chapter_name, course_name)
            elif get_worksheet_pdf_template:
                html_content = get_worksheet_pdf_template(html_content, chapter_name, course_name)
        
        # Save HTML temporarily
        temp_html = pdf_path.with_suffix('.temp.html')
        with open(temp_html, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Convert HTML to PDF with options
        if PDFKIT_AVAILABLE:
            try:
                options = {
                    'page-size': 'A4',
                    'margin-top': '0.75in',
                    'margin-right': '0.75in',
                    'margin-bottom': '0.75in',
                    'margin-left': '0.75in',
                    'encoding': "UTF-8",
                    'no-outline': None,
                    'enable-local-file-access': None
                }
                pdfkit.from_file(str(temp_html), str(pdf_path), options=options)
                temp_html.unlink()  # Delete temp file
                return True
            except Exception as e:
                print(f"  Error with pdfkit: {e}")
                temp_html.unlink()
                return False
        elif WEASYPRINT_AVAILABLE:
            try:
                WeasyHTML(filename=str(temp_html)).write_pdf(
                    str(pdf_path),
                    optimize_images=True
                )
                temp_html.unlink()  # Delete temp file
                return True
            except Exception as e:
                print(f"  Error with weasyprint: {e}")
                temp_html.unlink()
                return False
        else:
            print(f"  Warning: No PDF library available. Skipping PDF generation for {docx_path}")
            temp_html.unlink()
            return False
            
    except Exception as e:
        print(f"  Error converting {docx_path} to PDF: {e}")
        import traceback
        traceback.print_exc()
        return False


def comment_text_between_headings(html_content: str) -> str:
    """
    Comment out text between Heading1 (chapter name) and first Heading2 (section).
    Uses regex-based approach for more reliable HTML manipulation.
    """
    # Find h1 tag
    h1_match = re.search(r'<h1[^>]*>.*?</h1>', html_content, re.IGNORECASE | re.DOTALL)
    if not h1_match:
        return html_content  # No h1 found, return as-is
    
    h1_end = h1_match.end()
    
    # Find first h2 tag after h1
    remaining = html_content[h1_end:]
    h2_match = re.search(r'<h2[^>]*>', remaining, re.IGNORECASE)
    if not h2_match:
        return html_content  # No h2 found, return as-is
    
    h2_start = h1_end + h2_match.start()
    
    # Extract content between h1 and h2
    content_between = html_content[h1_end:h2_start]
    
    # Comment out the content (preserve whitespace)
    if content_between.strip():
        # Split by lines and comment each non-empty line
        lines = content_between.split('\n')
        commented_lines = []
        for line in lines:
            if line.strip() and not line.strip().startswith('<!--'):
                commented_lines.append(f"<!-- {line} -->")
            else:
                commented_lines.append(line)
        commented_content = '\n'.join(commented_lines)
        
        # Replace in original HTML
        return html_content[:h1_end] + commented_content + html_content[h2_start:]
    
    return html_content


def process_chapter(chapter_dir: Path, course_name: str, chapter_num: int, output_base: Path, 
                   force: bool = False, mcq_only: bool = False, html_only: bool = False,
                   chapter_name: Optional[str] = None):
    """
    Process a single chapter directory.
    
    Args:
        chapter_dir: Directory containing chapter files
        course_name: Name of the course
        chapter_num: Chapter number
        output_base: Base output directory
        force: Force regeneration of existing files
        mcq_only: Only process MCQs (extract, generate HTML, JSON, PDF)
        html_only: Only regenerate HTML files (chapter, intro)
    """
    print(f"  Processing chapter {chapter_num}...")
    
    # Find chapter docx file
    chapter_docx = None
    for file in chapter_dir.glob("chapter*.docx"):
        if not file.name.startswith("~$"):
            chapter_docx = file
            break
    
    if not chapter_docx:
        print(f"  Warning: No chapter*.docx found in {chapter_dir}")
        return None
    
    # Find intro and worksheet
    intro_docx = chapter_dir / "intro.docx"
    worksheet_docx = chapter_dir / "worksheet.docx"
    
    # Create output directory
    chapter_output = output_base / f"chapter_{chapter_num}"
    chapter_output.mkdir(parents=True, exist_ok=True)
    
    # Use provided chapter_name or default
    if not chapter_name:
        chapter_name = f"Chapter {chapter_num}"
    
    # Process chapter document
    chapter_name = f"Chapter {chapter_num}"
    
    # If HTML only mode, skip DOCX processing
    if html_only:
        # Only regenerate HTML files
        try:
            # Regenerate chapter HTML
            # Try to find modified chapter first, then original
            modified_chapter_path = chapter_output / f"chapter_{chapter_num}_modified.docx"
            if not modified_chapter_path.exists():
                modified_chapter_path = chapter_output / f"chapter{chapter_num}.docx"
            
            # If still not found, try source chapter
            if not modified_chapter_path.exists():
                modified_chapter_path = chapter_docx
            
            if modified_chapter_path.exists():
                chapter_html = convert_docx_to_html(modified_chapter_path)
                if chapter_html:
                    # Comment out text between Heading1 and first Heading2
                    chapter_html = comment_text_between_headings(chapter_html)
                    
                    chapter_html_path = chapter_output / f"chapter{chapter_num}.html"
                    if force or not chapter_html_path.exists():
                        with open(chapter_html_path, 'w', encoding='utf-8') as f:
                            f.write(chapter_html)
                        print(f"    Generated chapter HTML: {chapter_html_path.name}")
                    else:
                        print(f"    Chapter HTML already exists: {chapter_html_path.name} (use --force to regenerate)")
                else:
                    print(f"    Warning: Could not generate chapter HTML from {modified_chapter_path.name}")
            else:
                print(f"    Warning: No chapter DOCX found for HTML regeneration")
            
            # Regenerate intro HTML
            if intro_docx.exists():
                intro_html = convert_docx_to_html(intro_docx)
                if intro_html:
                    intro_html_path = chapter_output / "intro.html"
                    if force or not intro_html_path.exists():
                        with open(intro_html_path, 'w', encoding='utf-8') as f:
                            f.write(intro_html)
                        print(f"    Generated intro HTML: {intro_html_path.name}")
                    else:
                        print(f"    Intro HTML already exists: {intro_html_path.name} (use --force to regenerate)")
                else:
                    print(f"    Warning: Could not generate intro HTML")
            else:
                print(f"    No intro.docx found")
        except Exception as e:
            print(f"  Error regenerating HTML: {e}")
            import traceback
            traceback.print_exc()
        
        return chapter_output
    
    # If MCQ only mode, only process MCQs
    if mcq_only:
        try:
            chapter_doc = Document(str(chapter_docx))
            
            # Extract MCQs
            mcq_doc = extract_mcqs_from_chapter(chapter_doc)
            if not mcq_doc:
                print(f"    No MCQs found in chapter")
                return chapter_output
            
            # Process MCQs (same as below)
            # This will be handled in the MCQ processing section
        except Exception as e:
            print(f"  Error processing chapter document: {e}")
            import traceback
            traceback.print_exc()
            return chapter_output
    
    try:
        chapter_doc = Document(str(chapter_docx))
        
        # Extract MCQs
        mcq_doc = extract_mcqs_from_chapter(chapter_doc)
        if mcq_doc:
            # Save MCQ document (using chapter[no]-mcq.docx format)
            mcq_docx_path = chapter_output / f"chapter{chapter_num}-mcq.docx"
            if force or not mcq_docx_path.exists():
            mcq_doc.save(str(mcq_docx_path))
            print(f"    Extracted MCQs to {mcq_docx_path.name}")
            else:
                print(f"    MCQ DOCX already exists: {mcq_docx_path.name} (use --force to regenerate)")
            
            # Parse MCQ from DOCX to JSON first (needed for structured data)
            mcq_json_path = chapter_output / f"chapter{chapter_num}-mcq.json"
            mcq_data = None
            
            # Generate temporary HTML from DOCX for parsing
            temp_mcq_html_path = chapter_output / f"chapter{chapter_num}-mcq-temp.html"
            temp_mcq_html = None
            if force or not mcq_json_path.exists():
                # Generate temporary HTML from DOCX for parsing
                temp_mcq_html = convert_docx_to_html(mcq_docx_path)
                if temp_mcq_html:
                    with open(temp_mcq_html_path, 'w', encoding='utf-8') as f:
                        f.write(temp_mcq_html)
                    
                    # Parse MCQ from temporary HTML
                    try:
                        sys.path.insert(0, str(Path(__file__).parent))
                        from parse_mcq import parse_mcq_file
                        import json
                        
                        mcq_data = parse_mcq_file(temp_mcq_html_path)
                        if mcq_data and mcq_data.get('questions'):
                            # Track questions without correct answers
                            for q in mcq_data.get('questions', []):
                                if not q.get('correct_answer'):
                                    MISSING_ANSWERS_TRACKER.append({
                                        'course': course_name,
                                        'chapter': chapter_name,
                                        'chapter_num': chapter_num,
                                        'question_number': q.get('question_number', 0),
                                        'question_text': q.get('question_text', '')[:100]  # First 100 chars
                                    })
                            
                            with open(mcq_json_path, 'w', encoding='utf-8') as f:
                                json.dump(mcq_data, f, indent=2, ensure_ascii=False)
                            print(f"    Generated MCQ JSON: {mcq_json_path.name} ({mcq_data.get('total_questions', 0)} questions)")
                        else:
                            print(f"    Warning: No questions found in MCQ")
                    except Exception as e:
                        print(f"    Warning: Could not parse MCQ: {e}")
                        import traceback
                        traceback.print_exc()
                    
                    # Clean up temporary HTML
                    if temp_mcq_html_path.exists():
                        temp_mcq_html_path.unlink()
                else:
                    print(f"    Warning: Could not generate temporary MCQ HTML for parsing")
            else:
                # Load existing JSON
                if mcq_json_path.exists():
                    import json
                    with open(mcq_json_path, 'r', encoding='utf-8') as f:
                        mcq_data = json.load(f)
                    
                    # Track questions without correct answers from existing JSON
                    for q in mcq_data.get('questions', []):
                        if not q.get('correct_answer'):
                            MISSING_ANSWERS_TRACKER.append({
                                'course': course_name,
                                'chapter': chapter_name,
                                'chapter_num': chapter_num,
                                'question_number': q.get('question_number', 0),
                                'question_text': q.get('question_text', '')[:100]  # First 100 chars
                            })
                    
                    print(f"    Using existing MCQ JSON: {mcq_json_path.name} ({mcq_data.get('total_questions', 0)} questions)")
            
            # Generate MCQ HTML from JSON (primary method - more reliable)
            mcq_html_path = chapter_output / f"chapter{chapter_num}-mcq.html"
            if mcq_data and mcq_data.get('questions'):
                try:
                    from mcq_html_generator import generate_mcq_html_from_json
                    
                    # Generate HTML from JSON
                    mcq_html_from_json = generate_mcq_html_from_json(mcq_data)
                    
                    # Save HTML generated from JSON
                    if force or not mcq_html_path.exists():
                        with open(mcq_html_path, 'w', encoding='utf-8') as f:
                            f.write(mcq_html_from_json)
                        print(f"    Generated MCQ HTML from JSON: {mcq_html_path.name}")
                    else:
                        # Update HTML from JSON when force is used
                        if force:
                with open(mcq_html_path, 'w', encoding='utf-8') as f:
                                f.write(mcq_html_from_json)
                            print(f"    Regenerated MCQ HTML from JSON: {mcq_html_path.name}")
                        else:
                            print(f"    MCQ HTML already exists: {mcq_html_path.name} (use --force to regenerate)")
                except Exception as e:
                    print(f"    Warning: Could not generate MCQ HTML from JSON: {e}")
                    import traceback
                    traceback.print_exc()
            else:
                if not mcq_data:
                    print(f"    Warning: No MCQ data available, cannot generate HTML from JSON")
                elif not mcq_data.get('questions'):
                    print(f"    Warning: No questions in MCQ data, cannot generate HTML from JSON")
            
            # Generate MCQ PDF from JSON with professional formatting
            mcq_pdf_path = chapter_output / f"chapter{chapter_num}-mcq.pdf"
            
            if force or not mcq_pdf_path.exists():
                # Generate PDF from JSON instead of DOCX
                if mcq_data and mcq_json_path.exists():
                    try:
                        import json
                        from pdf_templates import get_mcq_pdf_from_json
                        
                        # Use already loaded mcq_data or load from file
                        if not mcq_data:
                            with open(mcq_json_path, 'r', encoding='utf-8') as f:
                                mcq_data = json.load(f)
                        
                        # Generate HTML from JSON
                        from pdf_templates import verify_mcq_pdf_completeness
                        mcq_html = get_mcq_pdf_from_json(mcq_data, chapter_name, course_name)
                        
                        # Verify completeness before generating PDF
                        verification = verify_mcq_pdf_completeness(mcq_data, mcq_html)
                        if not verification['all_match']:
                            print(f"    Warning: MCQ PDF verification issues:")
                            if not verification['questions_match']:
                                print(f"      Questions mismatch: JSON has {verification['total_questions']}, HTML has {verification['html_questions']}")
                            if not verification['options_match']:
                                print(f"      Options mismatch: JSON has {verification['total_options']}, HTML has {verification['html_options']}")
                            if not verification['answers_match']:
                                print(f"      Answers mismatch: JSON has {verification['total_answers']}, HTML has {verification['html_answers']}")
                            if verification['missing_items']:
                                for item in verification['missing_items'][:5]:  # Show first 5
                                    print(f"      Missing: {item}")
                        else:
                            print(f"    Verified: All {verification['total_questions']} questions, {verification['total_options']} options, {verification['total_answers']} answers included")
                        
                        # Save HTML temporarily
                        temp_html = mcq_pdf_path.with_suffix('.temp.html')
                        with open(temp_html, 'w', encoding='utf-8') as f:
                            f.write(mcq_html)
                        
                        # Convert HTML to PDF
                        if PDFKIT_AVAILABLE:
                            try:
                                options = {
                                    'page-size': 'A4',
                                    'margin-top': '0.75in',
                                    'margin-right': '0.75in',
                                    'margin-bottom': '0.75in',
                                    'margin-left': '0.75in',
                                    'encoding': "UTF-8",
                                    'no-outline': None,
                                    'enable-local-file-access': None
                                }
                                pdfkit.from_file(str(temp_html), str(mcq_pdf_path), options=options)
                                temp_html.unlink()
                                print(f"    Generated MCQ PDF from JSON: {mcq_pdf_path.name}")
                            except Exception as e:
                                print(f"    Error with pdfkit: {e}")
                                temp_html.unlink()
                        elif WEASYPRINT_AVAILABLE:
                            try:
                                WeasyHTML(filename=str(temp_html)).write_pdf(
                                    str(mcq_pdf_path),
                                    optimize_images=True
                                )
                                temp_html.unlink()
                                print(f"    Generated MCQ PDF from JSON: {mcq_pdf_path.name}")
                            except Exception as e:
                                print(f"    Error with weasyprint: {e}")
                                temp_html.unlink()
                        else:
                            print(f"    Warning: No PDF library available. Skipping MCQ PDF generation")
                            temp_html.unlink()
                    except Exception as e:
                        print(f"    Error generating MCQ PDF from JSON: {e}")
                        import traceback
                        traceback.print_exc()
                else:
                    print(f"    Warning: MCQ JSON not found, cannot generate PDF from JSON")
            else:
                print(f"    MCQ PDF already exists: {mcq_pdf_path.name} (use --force to regenerate)")
        
        # If MCQ only mode, skip chapter processing
        if not mcq_only:
        # Remove MCQs from chapter
        chapter_doc = remove_mcqs_from_chapter(chapter_doc)
        
        # Mark headings
        mark_headings_in_chapter(chapter_doc, chapter_name)
        
            # Save modified chapter (keep original name for reference)
        modified_chapter_path = chapter_output / f"chapter_{chapter_num}_modified.docx"
        chapter_doc.save(str(modified_chapter_path))
            
            # Also save with original name for consistency
            chapter_docx_original = chapter_output / f"chapter{chapter_num}.docx"
            chapter_doc.save(str(chapter_docx_original))
        
        # Generate chapter HTML
        chapter_html = convert_docx_to_html(modified_chapter_path)
        if chapter_html:
            # Comment out text between Heading1 and first Heading2
            chapter_html = comment_text_between_headings(chapter_html)
            
                chapter_html_path = chapter_output / f"chapter{chapter_num}.html"
                if force or not chapter_html_path.exists():
            with open(chapter_html_path, 'w', encoding='utf-8') as f:
                f.write(chapter_html)
            print(f"    Generated chapter HTML: {chapter_html_path.name}")
                else:
                    print(f"    Chapter HTML already exists: {chapter_html_path.name} (use --force to regenerate)")
        
    except Exception as e:
        print(f"  Error processing chapter document: {e}")
        import traceback
        traceback.print_exc()
    
    # Process intro (skip in MCQ only mode)
    if not mcq_only and intro_docx.exists():
        try:
            intro_html = convert_docx_to_html(intro_docx)
            if intro_html:
                intro_html_path = chapter_output / "intro.html"
                if force or not intro_html_path.exists():
                with open(intro_html_path, 'w', encoding='utf-8') as f:
                    f.write(intro_html)
                print(f"    Generated intro HTML: {intro_html_path.name}")
                else:
                    print(f"    Intro HTML already exists: {intro_html_path.name} (use --force to regenerate)")
        except Exception as e:
            print(f"  Error processing intro: {e}")
    
    # Process worksheet (skip in MCQ only mode)
    if not mcq_only and worksheet_docx.exists():
        try:
            worksheet_pdf_path = chapter_output / f"worksheet{chapter_num}.pdf"
            if force or not worksheet_pdf_path.exists():
                if docx_to_pdf(worksheet_docx, worksheet_pdf_path, pdf_type='worksheet',
                              chapter_name=chapter_name, course_name=course_name):
                print(f"    Generated worksheet PDF: {worksheet_pdf_path.name}")
            else:
                print(f"    Worksheet PDF already exists: {worksheet_pdf_path.name} (use --force to regenerate)")
        except Exception as e:
            print(f"  Error processing worksheet: {e}")
    
    return chapter_output


def process_course(course_dir: Path, output_base: Path, force: bool = False, 
                  mcq_only: bool = False, html_only: bool = False):
    """
    Process a single course directory.
    
    Args:
        course_dir: Directory containing course files
        output_base: Base output directory
        force: Force regeneration of existing files
        mcq_only: Only process MCQs
        html_only: Only regenerate HTML files
    """
    course_name = course_dir.name
    print(f"\nProcessing course: {course_name}")
    
    # Create output directory for course
    course_output = output_base / course_name
    course_output.mkdir(parents=True, exist_ok=True)
    
    # Process course intro and index (skip in MCQ only mode)
    if not mcq_only:
        # Read intro from SOURCE_DIR (original location)
        source_course_dir = SOURCE_DIR / course_name
        intro_course_docx = None
        intro_paths = [
            source_course_dir / "intro- course.docx",  # With space after hyphen
            source_course_dir / "intro-course.docx",
            source_course_dir / "course intro.docx",  # Alternative naming with space
            source_course_dir / "course_intro.docx",  # Alternative naming with underscore
        ]
        
        for path in intro_paths:
            if path.exists():
                intro_course_docx = path
                break
        
        if intro_course_docx and intro_course_docx.exists():
        try:
            intro_html = convert_docx_to_html(intro_course_docx)
            if intro_html:
                intro_path = course_output / "course_intro.html"
                    if force or not intro_path.exists():
                with open(intro_path, 'w', encoding='utf-8') as f:
                    f.write(intro_html)
                        print(f"  Generated course intro HTML from {intro_course_docx.name}")
                    else:
                        print(f"  Course intro HTML already exists (use --force to regenerate)")
                else:
                    print(f"  Warning: Course intro HTML is empty for {intro_course_docx.name}")
        except Exception as e:
            print(f"  Error processing course intro: {e}")
                import traceback
                traceback.print_exc()
        else:
            print(f"  Warning: Course intro DOCX not found in source directory")
            print(f"    Tried: {[p.name for p in intro_paths]}")
        
        # Process index (also from source directory)
        index_docx = source_course_dir / "index.docx"
    
    if index_docx.exists():
        try:
            index_html = convert_docx_to_html(index_docx)
            if index_html:
                index_path = course_output / "course_index.html"
                    if force or not index_path.exists():
                with open(index_path, 'w', encoding='utf-8') as f:
                    f.write(index_html)
                print(f"  Generated course index HTML")
                    else:
                        print(f"  Course index HTML already exists (use --force to regenerate)")
                else:
                    print(f"  Warning: Course index HTML is empty")
        except Exception as e:
            print(f"  Error processing course index: {e}")
        else:
            print(f"  Warning: Course index DOCX not found (index.docx)")
    
    # Process chapters
    chapter_dirs = sorted([d for d in course_dir.iterdir() if d.is_dir() and d.name.startswith("chapter")],
                         key=lambda x: int(re.search(r'\d+', x.name).group()) if re.search(r'\d+', x.name) else 0)
    
    # First, collect chapter names from source DOCX files
    source_course_dir = SOURCE_DIR / course_name
    chapter_names_map = {}  # Map chapter_num -> actual chapter name from DOCX
    
    for chapter_dir in chapter_dirs:
        chapter_num_match = re.search(r'\d+', chapter_dir.name)
        if chapter_num_match:
            chapter_num = int(chapter_num_match.group())
            
            # Get chapter name from source DOCX file
            source_chapter_dir = source_course_dir / chapter_dir.name
            chapter_docx_paths = [
                source_chapter_dir / f"chapter {chapter_num}.docx",
                source_chapter_dir / f"chapter_{chapter_num}.docx",
                source_chapter_dir / f"chapter{chapter_num}.docx",
            ]
            
            chapter_name_from_docx = None
            for docx_path in chapter_docx_paths:
                if docx_path.exists():
                    chapter_name_from_docx = extract_chapter_name_from_docx(docx_path)
                    if chapter_name_from_docx:
                        chapter_names_map[chapter_num] = chapter_name_from_docx
                        print(f"  Chapter {chapter_num} name from DOCX: {chapter_name_from_docx}")
                        break
            
            if not chapter_name_from_docx:
                print(f"  Warning: Could not extract chapter name from DOCX for Chapter {chapter_num}")
    
    # Process chapters
    for chapter_dir in chapter_dirs:
        chapter_num_match = re.search(r'\d+', chapter_dir.name)
        if chapter_num_match:
            chapter_num = int(chapter_num_match.group())
            # Use actual chapter name from DOCX if available
            actual_chapter_name = chapter_names_map.get(chapter_num, f"Chapter {chapter_num}")
            process_chapter(chapter_dir, course_name, chapter_num, course_output, 
                          force=force, mcq_only=mcq_only, html_only=html_only,
                          chapter_name=actual_chapter_name)
    
    return course_output


def main():
    """
    Main processing function.
    """
    global MISSING_ANSWERS_TRACKER
    # Reset tracker at start of each run
    MISSING_ANSWERS_TRACKER = []
    
    import argparse
    
    parser = argparse.ArgumentParser(description='Process Skill Lab Courses from DOCX files')
    parser.add_argument('--force', action='store_true',
                       help='Force regeneration of existing files (HTML, PDF, JSON)')
    parser.add_argument('--course', type=str,
                       help='Process only specific course (by name)')
    parser.add_argument('--mcq-only', action='store_true',
                       help='Only process MCQs (extract, generate HTML, JSON, PDF)')
    parser.add_argument('--html-only', action='store_true',
                       help='Only regenerate HTML files (chapter, intro)')
    args = parser.parse_args()
    
    # Validate arguments
    if args.mcq_only and args.html_only:
        print("Error: Cannot use --mcq-only and --html-only together")
        sys.exit(1)
    
    if not SOURCE_DIR.exists():
        print(f"Error: Source directory not found: {SOURCE_DIR}")
        sys.exit(1)
    
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    SCRIPT_DIR.mkdir(parents=True, exist_ok=True)
    
    print("=" * 60)
    print("Skill Lab Courses Processing Script")
    print("=" * 60)
    print(f"Source: {SOURCE_DIR}")
    print(f"Output: {OUTPUT_DIR}")
    print(f"Force mode: {args.force}")
    if args.mcq_only:
        print(f"Mode: MCQ Only (extract, HTML, JSON, PDF)")
    elif args.html_only:
        print(f"Mode: HTML Only (regenerate HTML files)")
    else:
        print(f"Mode: Full Processing")
    print()
    
    # Get all course directories
    course_dirs = []
    if args.course:
        course_path = SOURCE_DIR / args.course
        if not course_path.exists():
            print(f"Error: Course directory not found: {course_path}")
            sys.exit(1)
        course_dirs = [course_path]
    else:
    course_dirs = [d for d in SOURCE_DIR.iterdir() if d.is_dir() and not d.name.startswith('.')]
    
    course_dirs.sort()
    
    print(f"Found {len(course_dirs)} course(s) to process\n")
    
    processed = 0
    errors = 0
    
    for course_dir in course_dirs:
        try:
            process_course(course_dir, OUTPUT_DIR, force=args.force, 
                          mcq_only=args.mcq_only, html_only=args.html_only)
            processed += 1
        except Exception as e:
            errors += 1
            print(f"\nError processing {course_dir.name}: {e}")
            import traceback
            traceback.print_exc()
    
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Processed: {processed}")
    print(f"Errors: {errors}")
    print(f"Output directory: {OUTPUT_DIR.resolve()}")
    if args.force:
        print("\nNote: Files were regenerated (--force mode)")
    
    # Display summary of questions without correct answers
    if MISSING_ANSWERS_TRACKER:
        print("\n" + "=" * 60)
        print("QUESTIONS WITHOUT CORRECT ANSWERS - MANUAL VERIFICATION REQUIRED")
        print("=" * 60)
        print(f"Total questions missing correct answers: {len(MISSING_ANSWERS_TRACKER)}\n")
        
        # Group by course and chapter
        from collections import defaultdict
        by_course_chapter = defaultdict(list)
        for item in MISSING_ANSWERS_TRACKER:
            key = (item['course'], item['chapter'], item['chapter_num'])
            by_course_chapter[key].append(item)
        
        # Sort by course name, then chapter number
        sorted_items = sorted(by_course_chapter.items(), key=lambda x: (x[0][0], x[0][2]))
        
        for (course, chapter, chapter_num), questions in sorted_items:
            print(f"Course: {course}")
            print(f"  Chapter: {chapter} (Chapter {chapter_num})")
            print(f"  Missing answers for {len(questions)} question(s):")
            for q in sorted(questions, key=lambda x: x['question_number']):
                print(f"    - Question {q['question_number']}: {q['question_text']}")
            print()
        
        print("=" * 60)
        print("Note: These questions show 'Note: No correct answer specified for this question' in PDFs")
        print("=" * 60)
    else:
        print("\n✓ All MCQ questions have correct answers specified")
    
    print()


if __name__ == "__main__":
    main()
