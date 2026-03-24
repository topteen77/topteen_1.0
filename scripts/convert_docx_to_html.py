#!/usr/bin/env python3
"""
DOCX → HTML Converter
- Soft Enter (Shift+Enter) → <br>
- Hard Enter (Enter) → <p> or <li>
- No double bullets
- Nested lists in tables
- No borders, no inline styles
- 100% crash-proof
"""

import os
import sys
import re
import html
import base64
import string
from pathlib import Path
from typing import List, Tuple, Optional

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


# ----------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------
def detect_and_linkify_urls(text: str) -> str:
    """Detect URLs in text and convert them to HTML links."""
    import re
    # URL pattern to match http/https URLs
    url_pattern = r'(https?://[^\s<>"{}|\\^`\[\]]+)'
    
    def replace_url(match):
        url = match.group(1)
        return f'<a href="{url}" target="_blank">{url}</a>'
    
    return re.sub(url_pattern, replace_url, text)

def to_html_entities(text: str) -> str:
    """Convert special characters to HTML entities using library-based approach.

    This helps downstream systems (e.g., databases with limited charset)
    by avoiding direct insertion of characters like ₹, →, smart quotes, etc.
    """
    if not text:
        return text
    
    import html
    import unicodedata
    
    # First escape basic HTML characters
    text = html.escape(text, quote=False)
    
    # Convert remaining Unicode characters to numeric HTML entities
    result = []
    for char in text:
        # Skip already escaped characters (basic HTML entities)
        if char.startswith('&') and char.endswith(';'):
            result.append(char)
        else:
            # Convert Unicode character to numeric HTML entity
            codepoint = ord(char)
            if codepoint > 127:  # Non-ASCII characters
                result.append(f'&#{codepoint};')
            else:
                result.append(char)
    
    return ''.join(result)
def roman(n: int) -> str:
    vals = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
    syms = ["M","CM","D","CD","C","XC","L","XL","X","IX","V","IV","I"]
    res = ""
    for i, v in enumerate(vals):
        count = n // v
        res += syms[i] * count
        n -= v * count
    return res


def get_font_size(run) -> float:
    try:
        return run.font.size.pt if run.font.size else 11.0
    except:
        return 11.0


# ----------------------------------------------------------------------
# SAFE: Extract exact list prefix
# ----------------------------------------------------------------------
def get_numbering_text(paragraph) -> str:
    try:
        p = paragraph._p
        numPr = p.find('.//w:numPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if numPr is None:
            return ""

        ilvl_elem = numPr.find('.//w:ilvl', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        numId_elem = numPr.find('.//w:numId', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if numId_elem is None:
            return ""

        level = int(ilvl_elem.get(qn('w:val')) or 0) if ilvl_elem is not None else 0
        num_id = numId_elem.get(qn('w:val'))
        if not num_id:
            return ""

        numbering_part = paragraph.part.numbering_part
        if not numbering_part or not hasattr(numbering_part, 'numbering_definitions'):
            return "• "

        defs = numbering_part.numbering_definitions
        if not hasattr(defs, '_numbering') or not hasattr(defs._numbering, 'num'):
            return "• "

        num = defs._numbering.num.get(num_id)
        if not num or not hasattr(num, 'abstractNumId'):
            return "• "

        abstractNum = defs._numbering.abstractNum.get(num.abstractNumId.val)
        if not abstractNum or level >= len(abstractNum.lvl):
            return "• "

        lvl = abstractNum.lvl[level]
        numFmt = getattr(lvl.numFmt, 'val', 'bullet') if lvl.numFmt else 'bullet'
        start = int(getattr(lvl.start, 'val', 1)) if lvl.start else 1

        lvlText = getattr(lvl.lvlText, 'val', '') if lvl.lvlText else ''
        if not lvlText:
            if numFmt == "bullet": return "• "
            if numFmt == "decimal": lvlText = "%1."
            elif numFmt in ["lowerLetter", "upperLetter"]: lvlText = "%1)"
            elif numFmt in ["lowerRoman", "upperRoman"]: lvlText = "%1."
            else: return "• "

        counter = start

        def replace(m):
            nonlocal counter
            idx = int(m.group(1)) - 1
            if idx != 0: return "?"
            if numFmt == "decimal": return str(counter)
            if numFmt == "lowerRoman": return roman(counter).lower()
            if numFmt == "upperRoman": return roman(counter)
            if numFmt == "lowerLetter": return string.ascii_lowercase[(counter-1)%26]
            if numFmt == "upperLetter": return string.ascii_uppercase[(counter-1)%26]
            return str(counter)

        text = re.sub(r'%(\d)', replace, lvlText)
        return text + " "

    except Exception:
        return "• "


# ----------------------------------------------------------------------
# SAFE: Get list level and type
# ----------------------------------------------------------------------
def get_list_level_and_type(paragraph) -> Tuple[int, str]:
    try:
        pPr = paragraph._p.pPr
        if pPr is None:
            return -1, ""

        numPr = pPr.find('.//w:numPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if numPr is None:
            return -1, ""

        ilvl = numPr.find('.//w:ilvl', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        level = int(ilvl.get(qn('w:val')) or 0) if ilvl is not None else 0

        numId = numPr.find('.//w:numId', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if numId is not None:
            num_id = numId.get(qn('w:val'))
            try:
                defs = paragraph.part.numbering_part.numbering_definitions
                if hasattr(defs, '_numbering') and hasattr(defs._numbering, 'num'):
                    num = defs._numbering.num.get(num_id)
                    if num and hasattr(num, 'abstractNumId'):
                        abstractNum = defs._numbering.abstractNum.get(num.abstractNumId.val)
                        if abstractNum and level < len(abstractNum.lvl):
                            fmt = getattr(abstractNum.lvl[level].numFmt, 'val', 'bullet')
                            if fmt in ["decimal", "lowerRoman", "upperRoman", "lowerLetter", "upperLetter"]:
                                return level, "ol"
            except:
                pass
        return level, "ul"
    except:
        return -1, ""


# ----------------------------------------------------------------------
# Heading level: outline level (OOXML) first, then Word/LibreOffice style name or style_id
# ----------------------------------------------------------------------
W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# Set DEBUG_HEADINGS=1 (or true/yes) to print H1/H2/H3 detection to console.
# Example: DEBUG_HEADINGS=1 python scripts/convert_docx_to_html.py ...
#          DEBUG_HEADINGS=1 python manage.py convert_entrance_test_prep_docx --source ... --output ...
DEBUG_HEADINGS = os.environ.get("DEBUG_HEADINGS", "").strip().lower() in ("1", "true", "yes")


def _outline_level_from_paragraph(paragraph) -> int:
    """Return 1-6 from w:outlineLvl if set (0-based in OOXML: 0=first level), else 0."""
    try:
        p = paragraph._p
        if p is None:
            return 0
        pPr = p.find(".//w:pPr", namespaces=W_NS)
        if pPr is None:
            return 0
        outline = pPr.find(".//w:outlineLvl", namespaces=W_NS)
        if outline is None:
            return 0
        val = outline.get(qn("w:val"))
        if val is None:
            return 0
        n = int(val)
        # outline 0 = first heading level (h1), 1 = h2, etc.
        return max(1, min(6, n + 1))
    except Exception:
        return 0


def _style_id_from_paragraph_xml(paragraph) -> str:
    """Return the applied style id from w:pPr/w:pStyle @w:val (e.g. 'Heading2')."""
    try:
        p = paragraph._p
        if p is None:
            return ""
        pPr = p.find(".//w:pPr", namespaces=W_NS)
        if pPr is None:
            return ""
        pStyle = pPr.find(".//w:pStyle", namespaces=W_NS)
        if pStyle is None:
            return ""
        return (pStyle.get(qn("w:val")) or "").strip()
    except Exception:
        return ""


def heading_level_from_style(paragraph) -> int:
    """Return 1-6 if paragraph is a heading (outline level or style), else 0."""
    text_preview = (paragraph.text or "").strip()[:60].replace("\n", " ")
    # 1) OOXML outline level (most reliable; used by Word and LibreOffice)
    level = _outline_level_from_paragraph(paragraph)
    if level >= 1:
        if DEBUG_HEADINGS:
            print(f"[H{level}] outlineLvl -> text: {text_preview!r}")
        return level
    # 2) Applied style id from XML (e.g. Heading1, Heading2, Heading_20_2 for "Heading 2")
    style_id_xml = _style_id_from_paragraph_xml(paragraph)
    if style_id_xml:
        # LibreOffice can use _20_ for space in style id
        decoded = style_id_xml.replace("_20_", " ").replace("_", " ")
        m = re.match(r"Heading\s*(\d)", decoded, re.IGNORECASE)
        if m:
            n = int(m.group(1))
            if 1 <= n <= 6:
                if DEBUG_HEADINGS:
                    print(f"[H{n}] pStyle (decoded) -> text: {text_preview!r}  (w:pStyle={style_id_xml!r})")
                return n
        if re.match(r"Heading(\d)", style_id_xml, re.IGNORECASE):
            n = int(re.search(r"\d", style_id_xml).group())
            if 1 <= n <= 6:
                if DEBUG_HEADINGS:
                    print(f"[H{n}] pStyle (match) -> text: {text_preview!r}  (w:pStyle={style_id_xml!r})")
                return n
        if re.match(r"Title", style_id_xml, re.IGNORECASE):
            if DEBUG_HEADINGS:
                print(f"[H1] pStyle Title -> text: {text_preview!r}")
            return 1
    # 3) Style name / style_id from python-docx (e.g. "Heading 2", "Heading 2")
    try:
        style = paragraph.style
        if style is not None:
            name = (getattr(style, "name", None) or "").strip()
            style_id = (getattr(style, "style_id", None) or "").strip()
            for s in (name, style_id):
                if not s:
                    continue
                m = re.match(r"Heading\s*(\d)", s, re.IGNORECASE)
                if m:
                    n = int(m.group(1))
                    if 1 <= n <= 6:
                        if DEBUG_HEADINGS:
                            print(f"[H{n}] style.name/id -> text: {text_preview!r}  (name={name!r}, style_id={style_id!r})")
                        return n
                if re.match(r"Heading(\d)", s, re.IGNORECASE):
                    n = int(re.search(r"\d", s).group())
                    if 1 <= n <= 6:
                        if DEBUG_HEADINGS:
                            print(f"[H{n}] style (HeadingN) -> text: {text_preview!r}  (name={name!r}, style_id={style_id!r})")
                        return n
                if s == "Title":
                    if DEBUG_HEADINGS:
                        print(f"[H1] style Title -> text: {text_preview!r}")
                    return 1
    except Exception as e:
        if DEBUG_HEADINGS and (style_id_xml or "Heading" in (paragraph.style.name if getattr(paragraph, "style", None) else "")):
            print(f"[?] style exception: {e}  text: {text_preview!r}")
    # Debug: paragraph had heading-like style/XML but we didn't match -> will become <p>
    if DEBUG_HEADINGS:
        sn = (getattr(paragraph.style, "name", None) or "") if getattr(paragraph, "style", None) else ""
        if ("heading" in (style_id_xml or "").lower() or "title" in (style_id_xml or "").lower() or
                "heading" in sn.lower() or "title" in sn.lower()):
            print(f"[NOT HEADING] -> <p>  text: {text_preview!r}  w:pStyle={style_id_xml!r}  style.name={sn!r}")
    return 0


def slug_for_heading(text: str, used_ids: set) -> str:
    """Generate a unique id slug from heading text for anchor links."""
    text = re.sub(r"<[^>]+>", "", text).strip()
    slug = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")[:80]
    if not slug:
        slug = "section"
    base = slug
    c = 1
    while slug in used_ids:
        slug = f"{base}-{c}"
        c += 1
    used_ids.add(slug)
    return slug


# ----------------------------------------------------------------------
# HTML Builder
# ----------------------------------------------------------------------
class HTMLBuilder:
    def __init__(self, filename: str):
        self.lines: List[str] = []
        self.filename = filename
        self.title_set = False
        self.h1_set = False
        self.list_stack: List[Tuple[str, int]] = []
        self.heading_ids: set = set()  # per-instance, for unique id slugs

    def add(self, html_snippet: str):
        self.lines.append(html_snippet)

    def close_all_lists(self):
        while self.list_stack:
            lt, _ = self.list_stack.pop()
            self.add(f"</{lt}>")

    def get_html(self) -> str:
        self.close_all_lists()
        body = "\n".join(self.lines)
        # Derive <title> from first <h1>, fallback to filename
        try:
            m = re.search(r"<h1>(.*?)</h1>", body, flags=re.IGNORECASE|re.DOTALL)
            title_text = m.group(1).strip() if m else html.escape(self.filename, quote=False)
        except Exception:
            title_text = html.escape(self.filename, quote=False)

        # Minimal HTML wrapper with charset and title
        return (
            "<!DOCTYPE html>\n"
            "<html>\n"
            "<head>\n"
            "<meta charset=\"utf-8\">\n"
            f"<title>{title_text}</title>\n"
            "</head>\n"
            "<body>\n"
            f"{body}\n"
            "</body>\n"
            "</html>\n"
        )

    def ensure_list(self, level: int, list_type: str):
        while self.list_stack and self.list_stack[-1][1] > level:
            lt, _ = self.list_stack.pop()
            self.add(f"</{lt}>")
        while len(self.list_stack) <= level:
            self.add(f"<{list_type}>")
            self.list_stack.append((list_type, level))
        if self.list_stack and self.list_stack[-1][0] != list_type:
            lt, _ = self.list_stack.pop()
            self.add(f"</{lt}>")
            self.add(f"<{list_type}>")
            self.list_stack.append((list_type, level))

    def heading_level_from_size(self, size_pt: float) -> int:
        if size_pt >= 28: return 1
        if size_pt >= 22: return 2
        if size_pt >= 18: return 3
        if size_pt >= 15: return 4
        if size_pt >= 13: return 5
        return 6

    def maybe_make_heading(self, text: str, size_pt: float) -> Tuple[str, int]:
        level = self.heading_level_from_size(size_pt)
        if level == 1 and not self.h1_set:
            self.h1_set = True
            sid = slug_for_heading(text, self.heading_ids)
            return f'<h1 id="{sid}">{text}</h1>', 1
        if level == 1:
            level = 2
        sid = slug_for_heading(text, self.heading_ids)
        return f'<h{level} id="{sid}">{text}</h{level}>', level


# ----------------------------------------------------------------------
# Runs → HTML: SOFT ENTER → <br>
# ----------------------------------------------------------------------
def runs_to_html(runs) -> str:
    parts = []
    for run in runs:
        try:
            text = run.text
            # SOFT ENTER: \v → <br>
            if "\v" in text:
                segments = text.split("\v")
                for i, seg in enumerate(segments):
                    if i > 0:
                        parts.append("<br>")
                    parts.append(_format_run(seg, run))
                continue
            parts.append(_format_run(text, run))
        except Exception as e:
            print(f"Run processing error: {e}")
            continue
    return "".join(parts)


def _format_run(text: str, run) -> str:
    if not text:
        return ""
    text = text.replace("\t", "    ")
    
    # Handle explicit hyperlinks FIRST before HTML escaping
    try:
        if hasattr(run, 'hyperlink') and run.hyperlink:
            href = run.hyperlink.address or "#"
            if href and href != "#":
                # Don't escape the href URL, but escape the text content
                escaped_text = html.escape(text, quote=False)
                escaped_text = to_html_entities(escaped_text)
                return f'<a href="{href}" target="_blank">{escaped_text}</a>'
    except Exception as e:
        print(f"Hyperlink error: {e}")
        pass
    
    # Apply HTML escaping and entity conversion
    text = html.escape(text, quote=False)
    text = to_html_entities(text)
    
    # Detect and convert URLs in text to links
    text = detect_and_linkify_urls(text)

    try:
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"
    except:
        pass
    return text


# ----------------------------------------------------------------------
# Process paragraph – SOFT/HARD ENTER HANDLED
# ----------------------------------------------------------------------
def process_paragraph(builder: HTMLBuilder, paragraph) -> None:
    if not paragraph.text.strip() and not paragraph.runs:
        builder.add("<p>&nbsp;</p>")
        return

    # Prefer Word "Heading 1" / "Heading 2" / "Heading 3" styles for h1/h2/h3 (sidebar navigation)
    style_level = heading_level_from_style(paragraph)
    if style_level >= 1:
        builder.close_all_lists()
        runs_html = runs_to_html(paragraph.runs)
        if not runs_html.strip() and paragraph.text.strip():
            text = paragraph.text.strip()
            text = html.escape(text, quote=False)
            text = to_html_entities(text)
            runs_html = text
        if runs_html.strip():
            sid = slug_for_heading(runs_html, builder.heading_ids)
            if style_level == 1:
                builder.h1_set = True
            builder.add(f'<h{style_level} id="{sid}">{runs_html}</h{style_level}>')
        return

    try:
        level, list_type = get_list_level_and_type(paragraph)
        if level >= 0:
            prefix = get_numbering_text(paragraph).strip()
            item_html = runs_to_html(paragraph.runs)  # <br> for soft enter

            # Remove prefix if already in text - improved logic
            clean_html = item_html
            if prefix:
                # Check if the HTML content already starts with the prefix
                # Strip HTML tags temporarily to check text content
                text_content = re.sub(r'<[^>]+>', '', item_html).strip()
                if text_content.startswith(prefix.strip()):
                    # Remove the prefix from the HTML content
                    # Find where the prefix ends in the HTML
                    prefix_pattern = re.escape(prefix.strip())
                    clean_html = re.sub(f'^{prefix_pattern}\\s*', '', item_html, flags=re.IGNORECASE)
                    clean_html = re.sub(r'^[\s\u200B]+', '', clean_html)
                
                # For bullet lists, don't add the bullet prefix since <li> already provides it
                # For numbered lists, add the number prefix
                if list_type == "ul":
                    # Bullet list - don't add bullet prefix, just use clean content
                    # Also remove any existing bullet characters from the content
                    clean_html = re.sub(r'^[•·▪▫‣⁃\-\*\+]\s*', '', clean_html)
                    li_content = clean_html
                else:
                    # Numbered list - add the number prefix
                    li_content = f'{to_html_entities(html.escape(prefix))}{clean_html}'
            else:
                li_content = clean_html

            builder.ensure_list(level, list_type)
            builder.add(f"<li>{li_content}</li>")
            return
    except:
        pass

    # Check if this looks like a bullet point (starts with common bullet patterns)
    text = paragraph.text.strip()
    if text and not text.endswith(':'):
        # Check for common bullet point patterns
        bullet_patterns = [
            r'^[•·▪▫‣⁃]',  # Unicode bullets
            r'^[-*+]',      # ASCII bullets
            r'^\d+[\.\)]',  # Numbered lists
            r'^[a-zA-Z][\.\)]',  # Letter lists
        ]
        
        for pattern in bullet_patterns:
            if re.match(pattern, text):
                # This looks like a bullet point, treat as list item
                builder.ensure_list(0, "ul")
                # Remove bullet characters from the content since <li> provides styling
                content = runs_to_html(paragraph.runs)
                content = re.sub(r'^[•·▪▫‣⁃\-\*\+]\s*', '', content)
                builder.add(f"<li>{content}</li>")
                return

    # Close any open lists before starting a normal paragraph
    builder.close_all_lists()

    # Normal paragraph (soft enters → <br>)
    try:
        runs_html = runs_to_html(paragraph.runs)
        
        # Handle case where paragraph has text but no runs
        if not runs_html.strip() and paragraph.text.strip():
            # Process the paragraph text directly
            text = paragraph.text.strip()
            text = html.escape(text, quote=False)
            text = to_html_entities(text)
            text = detect_and_linkify_urls(text)
            runs_html = text
        
        if not runs_html.strip():
            builder.add("<p>&nbsp;</p>")
            return

        max_size = max((get_font_size(r) for r in paragraph.runs), default=11)
        if max_size >= 13:
            heading_tag, lvl = builder.maybe_make_heading(runs_html, max_size)
            if DEBUG_HEADINGS:
                text_preview = re.sub(r"<[^>]+>", "", runs_html).strip()[:60]
                print(f"[H{lvl}] size ({max_size}pt) -> text: {text_preview!r}")
            builder.add(heading_tag)
        else:
            builder.add(f"<p>{runs_html}</p>")
    except:
        builder.add("<p>&nbsp;</p>")


# ----------------------------------------------------------------------
# Process cell content
# ----------------------------------------------------------------------
def process_cell_content(builder: HTMLBuilder, cell) -> None:
    for para in cell.paragraphs:
        process_paragraph(builder, para)


# ----------------------------------------------------------------------
# Table → HTML
# ----------------------------------------------------------------------
def table_to_html(builder: HTMLBuilder, table) -> None:
    builder.add("<table>")
    for row_idx, row in enumerate(table.rows):
        builder.add("<tr>")
        for cell in row.cells:
            colspan = getattr(cell, "grid_span", 1)
            tag = "th" if row_idx == 0 else "td"
            attr = f' colspan="{colspan}"' if colspan > 1 else ""
            builder.add(f"<{tag}{attr}>")
            cell_builder = HTMLBuilder("")
            cell_builder.list_stack = builder.list_stack.copy()
            process_cell_content(cell_builder, cell)
            cell_builder.close_all_lists()
            builder.add("".join(cell_builder.lines))
            builder.add(f"</{tag}>")
        builder.add("</tr>")
    builder.add("</table>")


# ----------------------------------------------------------------------
# Image → data URI
# ----------------------------------------------------------------------
def image_to_data_uri(part, rel) -> str:
    try:
        image_bytes = rel.target_ref.blob
        mime = rel.target_part.content_type
        b64 = base64.b64encode(image_bytes).decode()
        return f'data:{mime};base64,{b64}'
    except:
        return ""


# ----------------------------------------------------------------------
# Main conversion
# ----------------------------------------------------------------------
def convert_docx_to_html(docx_path: Path) -> Optional[str]:
    try:
        if DEBUG_HEADINGS:
            print(f"\n--- DEBUG_HEADINGS: {docx_path.name} ---")
        doc = Document(str(docx_path))
        filename = docx_path.stem
        builder = HTMLBuilder(filename)

        for block in doc.element.body:
            if block.tag.endswith("p"):
                para_idx = [p._p for p in doc.paragraphs].index(block)
                process_paragraph(builder, doc.paragraphs[para_idx])
            elif block.tag.endswith("tbl"):
                tbl_idx = [t._tbl for t in doc.tables].index(block)
                table_to_html(builder, doc.tables[tbl_idx])

        html_out = builder.get_html()

        # Embed images
        for rel in doc.part.related_parts.values():
            if rel.content_type.startswith("image/"):
                placeholder = f'[IMAGE:{rel.partname}]'
                if placeholder in html_out:
                    data_uri = image_to_data_uri(doc.part, rel)
                    if data_uri:
                        img_tag = f'<img src="{data_uri}" alt="image">'
                        html_out = html_out.replace(placeholder, img_tag)

        return html_out

    except Exception as e:
        print(f"Error processing {docx_path}: {e}")
        return None


# ----------------------------------------------------------------------
# Directory processor
# ----------------------------------------------------------------------
def process_directory(source_dir: str, output_dir: str):
    src = Path(source_dir)
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    docx_files = [f for f in src.rglob("*.docx") if not f.name.startswith("~$")]
    print(f"Found {len(docx_files)} .docx files")

    success = error = 0
    for docx in docx_files:
        try:
            html_content = convert_docx_to_html(docx)
            if not html_content:
                error += 1
                continue

            rel = docx.relative_to(src)
            out_file = (out / rel).with_suffix(".txt")
            out_file.parent.mkdir(parents=True, exist_ok=True)

            with open(out_file, "w", encoding="utf-8") as f:
                f.write(html_content)

            success += 1
            print(f"Success: {rel}")

        except Exception as e:
            error += 1
            print(f"Failed: {docx}: {e}")

    print("\n=== SUMMARY ===")
    print(f"Processed : {success}")
    print(f"Errors    : {error}")
    print(f"Output    : {out.resolve()}")


# ----------------------------------------------------------------------
# CLI
# ----------------------------------------------------------------------
def main():
    source = "/home/itpc6/Public/share/content- Topteen/career library 2025/final careers"
    output = "career_html_output"

    if len(sys.argv) > 1:
        source = sys.argv[1]
    if len(sys.argv) > 2:
        output = sys.argv[2]

    src_path = Path(source)
    if not src_path.exists():
        print(f"Source not found: {source}")
        sys.exit(1)

    print("DOCX → HTML: Soft Enter → <br>, Hard Enter → <p>/<li>")
    print(f"Source : {source}")
    print(f"Output : {output}\n")

    # Single file mode: if source is a .docx file, convert it and write to output dir
    if src_path.is_file() and src_path.suffix.lower() == ".docx":
        out_dir = Path(output)
        out_dir.mkdir(parents=True, exist_ok=True)
        out_file = out_dir / src_path.with_suffix(".txt").name
        html_content = convert_docx_to_html(src_path)
        if html_content:
            out_file.write_text(html_content, encoding="utf-8")
            print(f"Success: {src_path.name} → {out_file}")
        else:
            print(f"Failed: {src_path.name}")
            sys.exit(1)
        return

    if not src_path.is_dir():
        print(f"Source is not a directory or .docx file: {source}")
        sys.exit(1)

    process_directory(source, output)


if __name__ == "__main__":
    main()