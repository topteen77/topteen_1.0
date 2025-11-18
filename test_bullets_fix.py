#!/usr/bin/env python3
"""
Test script to create a simple DOCX with bullet points and convert it to HTML
to verify the double bullets fix.
"""

from docx import Document
from docx.shared import Inches
import tempfile
import os
import sys

# Add the scripts directory to path
sys.path.append('/home/itpc6/Public/django/git-repo/25oct/demo-topteens/scripts')

def create_test_docx():
    """Create a test DOCX file with bullet points"""
    doc = Document()
    
    # Add a title
    doc.add_heading('Test Document with Bullets', 0)
    
    # Add a paragraph
    doc.add_paragraph('This document contains bullet points to test the conversion.')
    
    # Add bullet points
    doc.add_paragraph('First bullet point', style='List Bullet')
    doc.add_paragraph('Second bullet point', style='List Bullet')
    doc.add_paragraph('Third bullet point', style='List Bullet')
    
    # Add a numbered list
    doc.add_paragraph('First numbered item', style='List Number')
    doc.add_paragraph('Second numbered item', style='List Number')
    doc.add_paragraph('Third numbered item', style='List Number')
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(temp_file.name)
    return temp_file.name

def test_conversion():
    """Test the DOCX to HTML conversion"""
    # Create test DOCX
    docx_file = create_test_docx()
    print(f"Created test DOCX: {docx_file}")
    
    # Import the conversion function
    from convert_docx_to_html import convert_docx_to_html
    
    # Convert to HTML
    html_content = convert_docx_to_html(docx_file)
    
    # Save HTML output
    html_file = docx_file.replace('.docx', '.html')
    with open(html_file, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    print(f"Created HTML output: {html_file}")
    
    # Show the HTML content
    print("\n=== HTML OUTPUT ===")
    print(html_content)
    
    # Clean up
    os.unlink(docx_file)
    os.unlink(html_file)
    
    return html_content

if __name__ == "__main__":
    test_conversion()
