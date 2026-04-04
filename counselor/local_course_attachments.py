"""
Resolve lesson / index Word files under career-counsellor content folders for mindmap export.

Expected layout (e.g. …/career counsellor course/):
  attachments/career counselling course- Index.docx
  attachments/chapter N/chapter N part M.docx
  Quiz/chapter N/…docx
  videos/chapter N/…mp4
"""
from __future__ import annotations

import re
from pathlib import Path


def index_docx_path(content_root: Path) -> Path | None:
    p = content_root / "attachments" / "career counselling course- Index.docx"
    return p if p.is_file() else None


def chapter_part_lesson_docx(content_root: Path, chapter_ord: int, part_ord: int) -> Path | None:
    """Primary lesson handout for a chapter/part (ordinal positions, 1-based)."""
    att = content_root / "attachments" / f"chapter {chapter_ord}"
    if not att.is_dir():
        return None
    name = f"chapter {chapter_ord} part {part_ord}.docx"
    p = att / name
    if p.is_file():
        return p
    alt = att / name.replace("chapter ", "Chapter ", 1)
    return alt if alt.is_file() else None


def _paragraph_is_center_aligned(para) -> bool:
    """Document title / part banner lines are often center-aligned; omit from mindmap."""
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        al = para.paragraph_format.alignment
        return al == WD_ALIGN_PARAGRAPH.CENTER
    except Exception:  # pragma: no cover
        return False


def _paragraph_is_heading_style(para) -> bool:
    """Bold/heading paragraphs only (not body copy), for mindmap labels."""
    text = (para.text or "").strip()
    if not text or len(text) > 220:
        return False
    st = (getattr(para.style, "name", None) or "")
    if st.startswith("Heading") or "heading" in st.lower():
        return True
    sl = st.lower()
    if sl in ("title", "subtitle") or "title" in sl:
        return True
    runs = [r for r in para.runs if r.text]
    if not runs:
        return False
    bold_chars = sum(len(r.text) for r in runs if r.bold)
    total = sum(len(r.text) for r in runs)
    if total == 0:
        return False
    return (bold_chars / total) >= 0.85


def extract_mindmap_headings_from_docx(path: Path, *, max_items: int = 32) -> list[str]:
    """
    Short bold or Word-heading lines from a lesson handout (e.g. 'What Is Counselling?').
    Omits body paragraphs.
    """
    try:
        from docx import Document
    except ImportError:  # pragma: no cover
        return []
    try:
        doc = Document(str(path))
    except Exception:
        return []
    seen: set[str] = set()
    out: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        if _paragraph_is_center_aligned(para):
            continue
        if not _paragraph_is_heading_style(para):
            continue
        key = text.casefold()
        if key in seen:
            continue
        seen.add(key)
        out.append(text[:200])
        if len(out) >= max_items:
            break
    return out


def extract_plain_text_from_docx(path: Path, *, max_chars: int = 120_000) -> str:
    try:
        from docx import Document
    except ImportError:  # pragma: no cover
        return ""
    try:
        doc = Document(str(path))
    except Exception:
        return ""
    chunks: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            chunks.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    chunks.append(t)
    text = " ".join(chunks)
    if len(text) > max_chars:
        text = text[:max_chars]
    return text


def lesson_headings_for_part(
    part,
    *,
    use_remote_pdf: bool,
    base_dir: Path | None,
    content_root: Path | None,
    chapter_ord: int,
    part_ord: int,
    max_items: int = 24,
) -> list[str]:
    """
    Important headings from lesson handout (bold / Heading styles), else brief PDF-derived lines.
    """
    if content_root:
        p = chapter_part_lesson_docx(content_root, chapter_ord, part_ord)
        if p:
            h = extract_mindmap_headings_from_docx(p, max_items=max_items)
            if h:
                return h
    if use_remote_pdf:
        from counselor.pdf_mindmap_extract import extract_text_from_pdf_bytes, resolve_part_pdf_bytes, text_to_outline_bullets

        data = resolve_part_pdf_bytes(part, base_dir=base_dir)
        if data:
            t = extract_text_from_pdf_bytes(data)
            if t.strip():
                return text_to_outline_bullets(t, max_items=min(6, max_items))
    return []


def raw_index_headings(content_root: Path | None, *, max_items: int = 18) -> list[str]:
    if not content_root:
        return []
    p = index_docx_path(content_root)
    if not p:
        return []
    return extract_mindmap_headings_from_docx(p, max_items=max_items)


def _parse_chapter_part_from_stem(stem: str) -> tuple[int, int] | None:
    """
    Map filenames like 'chapter 5 part 3', 'Chapter 4 Part3', 'chapter 10 introduction' to (chapter, part).
    Odd names like 'part -1' map part to 1.
    """
    s = stem.lower().replace("_", " ")
    s = re.sub(r"\s+", " ", s).strip()
    if re.search(r"new microsoft|^\s*~\$", s):
        return None
    # Course-level intro (handled separately)
    if "course" in s and "intro" in s and "chapter" not in s:
        return None
    m = re.search(r"chapter\s*(\d+).*?intro", s)
    if m and "part" not in s:
        return (int(m.group(1)), 1)
    m = re.search(r"chapter\s*(\d+).*?part\s*(-?\d+)", s, re.I)
    if m:
        ch, pt = int(m.group(1)), int(m.group(2))
        if pt < 1:
            pt = 1
        return (ch, pt)
    m = re.match(r"chapter\s*(\d+)\s*$", s)
    if m:
        return (int(m.group(1)), 1)
    m = re.search(r"chapter\s*(\d+)", s)
    if m and "part" not in s and "intro" in s:
        return (int(m.group(1)), 1)
    return None


def scan_quiz_video_sets(
    content_root: Path | None,
) -> tuple[set[tuple[int, int]], set[tuple[int, int]], bool]:
    """
    (chapter, part) pairs that have a Quiz docx or Video mp4; True if course-level intro video exists.
    """
    quiz: set[tuple[int, int]] = set()
    video: set[tuple[int, int]] = set()
    course_intro_video = False
    if not content_root or not content_root.is_dir():
        return quiz, video, course_intro_video

    qroot = content_root / "Quiz"
    if qroot.is_dir():
        for f in qroot.rglob("*.docx"):
            pair = _parse_chapter_part_from_stem(f.stem)
            if pair:
                quiz.add(pair)

    vroot = content_root / "videos"
    if vroot.is_dir():
        for f in vroot.rglob("*.mp4"):
            rel = f.relative_to(vroot)
            if rel.parts and rel.parts[0].lower().startswith("course"):
                course_intro_video = True
                continue
            pair = _parse_chapter_part_from_stem(f.stem)
            if pair:
                video.add(pair)

    return quiz, video, course_intro_video
