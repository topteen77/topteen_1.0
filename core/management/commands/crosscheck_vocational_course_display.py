"""
Cross-check VocationalCourse.content_html (admin description) vs content_json
as used on the public accordion (overview + sections html + pros + cons).

Optional: compare plain text from Word (.docx) files under a folder tree to the
same course's content_html and JSON-assembled text (filename stem vs course name).

Reports per course:
- Missing JSON or very low text coverage vs source HTML
- Sections where stored JSON body is much smaller than a Python re-parse of the
  same HTML suggests (possible missing chunks; admin JS may differ slightly)
- Duplicate or near-duplicate body text across two section keys
- Repeated bullet titles inside pros / cons node HTML
- With --word-root: Word vs DB HTML / JSON alignment, missing Word file, orphan .docx
"""
from __future__ import annotations

import json
import re
from collections import Counter
from difflib import SequenceMatcher
from pathlib import Path

from bs4 import BeautifulSoup
from django.core.management.base import BaseCommand

from core.management.commands.fix_vocational_course_headings import (
    FIXED_HEADINGS,
    generate_content_json,
)
from core.models import VocationalCourse


def section_key(heading: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", heading.lower())


def plain_text(html: str | None) -> str:
    if not html or not str(html).strip():
        return ""
    soup = BeautifulSoup(str(html), "html.parser")
    return re.sub(r"\s+", " ", soup.get_text(" ", strip=True)).strip()


def stored_display_plain(content_json: dict | None) -> str:
    """Plain text approximately shown on frontend from JSON (overview + all sections)."""
    if not content_json or not isinstance(content_json, dict):
        return ""
    parts: list[str] = []
    parts.append(plain_text(content_json.get("overview")))
    sections = content_json.get("sections") or {}
    for heading in FIXED_HEADINGS:
        if heading == "Overview":
            continue
        key = section_key(heading)
        block = sections.get(key)
        if not block:
            continue
        if isinstance(block, str):
            parts.append(plain_text(block))
        elif isinstance(block, dict):
            parts.append(plain_text(block.get("html")))
            parts.append(plain_text(block.get("pros")))
            parts.append(plain_text(block.get("cons")))
    return " ".join(p for p in parts if p)


def similarity(a: str, b: str) -> float:
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()


def list_item_signatures(inner_html: str | None) -> list[str]:
    """First meaningful line per <li> for duplicate detection."""
    if not inner_html or not str(inner_html).strip():
        return []
    wrap = f"<ul>{inner_html}</ul>"
    soup = BeautifulSoup(wrap, "html.parser")
    sigs = []
    for li in soup.find_all("li"):
        t = plain_text(str(li))[:200]
        if len(t) > 3:
            sigs.append(t.lower())
    return sigs


def duplicate_li_titles(inner_html: str | None) -> list[str]:
    sigs = list_item_signatures(inner_html)
    counts = Counter(sigs)
    return [s for s, c in counts.items() if c > 1]


def normalize_for_match(name: str) -> str:
    """Normalize course name or file stem for matching .docx to VocationalCourse.name."""
    n = (name or "").lower().strip()
    n = n.replace(".docx", "").replace(".doc", "")
    n = re.sub(r"[\u2018\u2019]", "'", n)
    n = re.sub(r"[^\w\s&]", " ", n)
    n = re.sub(r"\s+", " ", n).strip()
    return n


def scan_docx_files(word_root: Path) -> dict[str, str]:
    """Map normalized stem -> absolute path (last wins if duplicates)."""
    out: dict[str, str] = {}
    if not word_root.is_dir():
        return out
    for p in word_root.rglob("*.docx"):
        if not p.is_file():
            continue
        if p.name.startswith("~$") or p.name.startswith(".~"):
            continue
        stem = normalize_for_match(p.stem)
        if stem:
            out[stem] = str(p.resolve())
    return out


def extract_docx_paragraphs(path: str) -> list[str]:
    """Non-empty paragraph and table cell texts in document order (for snippet-level QA)."""
    from docx import Document

    doc = Document(path)
    chunks: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            chunks.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    chunks.append(t)
    return chunks


def extract_docx_plain(path: str) -> str:
    """Single-line plain text from a .docx (paragraphs and table cells)."""
    inner = [re.sub(r"\s+", " ", p) for p in extract_docx_paragraphs(path)]
    return re.sub(r"\s+", " ", " ".join(inner)).strip()


def _normalize_match_line(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").lower()).strip()


_WORD_HEADING_BODY_SPLIT = re.compile(
    r"^(overview|conclusion|introduction|summary)\s*\n",
    re.I,
)


def _word_paragraph_plain_variants(raw: str) -> list[str]:
    """Whole paragraph plus body-after-known-heading-line (Word often merges heading + body)."""
    raw = (raw or "").strip()
    if not raw:
        return []
    variants = [raw]
    m = _WORD_HEADING_BODY_SPLIT.match(raw)
    if m:
        body = raw[m.end() :].strip()
        if len(body) >= 40:
            variants.append(body)
    return variants


def long_word_paragraphs_missing_from_html(
    paragraphs: list[str],
    html_plain: str,
    min_len: int = 52,
    max_report: int = 6,
) -> list[str]:
    """
    Flag long Word paragraphs whose normalized text does not appear in content_html plain text.
    Catches missing conclusion bodies, merged heading+body in Word vs empty <h2> in HTML, etc.
    """
    if not html_plain or not paragraphs:
        return []
    h = _normalize_match_line(html_plain)
    out: list[str] = []
    for para in paragraphs:
        raw = (para or "").strip()
        if len(raw) < min_len:
            continue
        if any(_normalize_match_line(v) in h for v in _word_paragraph_plain_variants(raw)):
            continue
        n = _normalize_match_line(raw)
        if not n:
            continue
        # Skip common boilerplate / very generic one-liners
        if re.fullmatch(r"conclusion|references?|bibliography", n):
            continue
        snippet = raw[:140] + ("…" if len(raw) > 140 else "")
        out.append(snippet)
        if len(out) >= max_report:
            break
    return out


def long_paragraphs_in_html_not_in_json(
    paragraphs: list[str],
    html_plain: str,
    json_plain: str,
    min_len: int = 52,
    max_report: int = 4,
) -> list[str]:
    """Word paragraphs that appear in DB HTML but not in JSON-assembled display (split pipeline gap)."""
    if not html_plain or not json_plain or not paragraphs:
        return []
    h = _normalize_match_line(html_plain)
    j = _normalize_match_line(json_plain)
    out: list[str] = []
    for para in paragraphs:
        raw = (para or "").strip()
        if len(raw) < min_len:
            continue
        in_html = any(_normalize_match_line(v) in h for v in _word_paragraph_plain_variants(raw))
        in_json = any(_normalize_match_line(v) in j for v in _word_paragraph_plain_variants(raw))
        if not in_html or in_json:
            continue
        out.append(raw[:140] + ("…" if len(raw) > 140 else ""))
        if len(out) >= max_report:
            break
    return out


def find_word_doc_path(course_name: str, stem_to_path: dict[str, str]) -> tuple[str | None, str | None]:
    """
    Return (path, match_note) where match_note is None for exact key match,
    or 'fuzzy:Other Stem' when matched by similarity.
    """
    key = normalize_for_match(course_name)
    if key in stem_to_path:
        return stem_to_path[key], None
    best_stem, best_r, best_path = "", 0.0, None
    for stem, path in stem_to_path.items():
        r = SequenceMatcher(None, key, stem).ratio()
        if r > best_r:
            best_r, best_stem, best_path = r, stem, path
    if best_path and best_r >= 0.86:
        return best_path, f"fuzzy match to file stem (similarity {best_r:.2f} vs «{best_stem}»)"
    return None, None


def word_doc_issues_for_course(
    course: VocationalCourse,
    doc_path: str,
    doc_plain: str,
    match_note: str | None,
) -> list[dict]:
    issues: list[dict] = []
    if match_note:
        issues.append({"type": "WORD_FILE_FUZZY_MATCH", "detail": match_note, "path": doc_path})

    html_plain = plain_text(course.content_html)
    json_plain = stored_display_plain(course.content_json)

    if len(doc_plain) > 200 and html_plain:
        cov = similarity(doc_plain, html_plain)
        char_ratio = len(html_plain) / max(len(doc_plain), 1)
        if cov < 0.62 and char_ratio < 0.72:
            issues.append(
                {
                    "type": "WORD_VS_HTML_DIVERGENT",
                    "detail": f"Word doc vs content_html: similarity {cov:.2f}, HTML length {char_ratio:.0%} of Word text. "
                    "Editor HTML may be missing or differ strongly from the Word source.",
                    "similarity": round(cov, 3),
                    "char_ratio": round(char_ratio, 3),
                    "word_chars": len(doc_plain),
                    "html_chars": len(html_plain),
                    "path": doc_path,
                }
            )

    if len(doc_plain) > 200 and json_plain:
        cov = similarity(doc_plain, json_plain)
        char_ratio = len(json_plain) / max(len(doc_plain), 1)
        if cov < 0.58 and char_ratio < 0.70:
            issues.append(
                {
                    "type": "WORD_VS_JSON_DIVERGENT",
                    "detail": f"Word doc vs JSON-assembled text: similarity {cov:.2f}, assembled {char_ratio:.0%} of Word length.",
                    "similarity": round(cov, 3),
                    "char_ratio": round(char_ratio, 3),
                    "path": doc_path,
                }
            )

    try:
        paras = extract_docx_paragraphs(doc_path)
    except Exception:
        paras = []

    if paras and html_plain:
        missing_html = long_word_paragraphs_missing_from_html(paras, html_plain)
        if missing_html:
            issues.append(
                {
                    "type": "WORD_PARAGRAPH_NOT_IN_HTML",
                    "detail": (
                        f"{len(missing_html)} long Word paragraph(s) not found verbatim in content_html "
                        "(import may have dropped text or headings). Examples: "
                        + " | ".join(missing_html)
                    ),
                    "path": doc_path,
                }
            )

    if paras and html_plain and json_plain:
        missing_json = long_paragraphs_in_html_not_in_json(paras, html_plain, json_plain)
        if missing_json:
            issues.append(
                {
                    "type": "WORD_IN_HTML_BUT_NOT_JSON_DISPLAY",
                    "detail": (
                        "Text present in content_html but not in JSON-assembled display (regenerate JSON). "
                        + " | ".join(missing_json)
                    ),
                    "path": doc_path,
                }
            )

    return issues


def analyze_course(course: VocationalCourse) -> dict:
    cid = course.pk
    name = course.name or ""
    html = (course.content_html or "").strip()
    stored = course.content_json

    issues: list[dict] = []

    if not html:
        return {
            "id": cid,
            "name": name,
            "issues": [{"type": "NO_SOURCE_HTML", "detail": "No content_html; nothing to compare."}],
        }

    if not stored or not isinstance(stored, dict):
        issues.append(
            {
                "type": "MISSING_JSON",
                "detail": "content_json is empty; public page uses HTML fallback, not accordion JSON.",
            }
        )
        return {"id": cid, "name": name, "issues": issues}

    src_plain = plain_text(html)
    disp_plain = stored_display_plain(stored)

    # Order-independent similarity can be low when paragraphs are reordered; require
    # both low similarity and noticeably fewer assembled characters than source.
    if len(src_plain) > 400:
        cov = similarity(src_plain, disp_plain)
        char_ratio = len(disp_plain) / max(len(src_plain), 1)
        if cov < 0.68 and char_ratio < 0.78:
            issues.append(
                {
                    "type": "LOW_TEXT_COVERAGE",
                    "detail": f"Similarity vs JSON-assembled text ≈ {cov:.2f}, assembled length ≈ {char_ratio:.0%} of source. "
                    "Substantial description text may be missing from overview/sections in content_json.",
                    "similarity": round(cov, 3),
                    "char_ratio": round(char_ratio, 3),
                    "source_chars": len(src_plain),
                    "assembled_chars": len(disp_plain),
                }
            )

    try:
        regen = generate_content_json(html, course_name=name)
    except Exception as e:
        issues.append({"type": "REGEN_ERROR", "detail": str(e)})
        regen = {"sections": {}}

    stored_sections = stored.get("sections") or {}
    regen_sections = regen.get("sections") or {}

    for heading in FIXED_HEADINGS:
        if heading == "Overview":
            continue
        key = section_key(heading)
        st = stored_sections.get(key) or {}
        rg = regen_sections.get(key) or {}
        if not isinstance(st, dict):
            st = {"html": str(st), "title": heading}
        if not isinstance(rg, dict):
            rg = {"html": str(rg), "title": heading}
        st_body = plain_text(st.get("html") or "")
        rg_body = plain_text(rg.get("html") or "")
        if len(rg_body) > 250 and len(st_body) < 40:
            issues.append(
                {
                    "type": "SECTION_EMPTY_IN_JSON_BUT_REGEN_HAS_BODY",
                    "section": key,
                    "detail": f"Section '{key}' is nearly empty in content_json but re-parse finds ~{len(rg_body)} chars "
                    f"from content_html. Regenerate JSON in admin (Generate Accordion from Content).",
                    "stored_chars": len(st_body),
                    "regen_chars": len(rg_body),
                }
            )
        elif len(rg_body) > 200 and len(st_body) >= 40 and len(st_body) < len(rg_body) * 0.5:
            issues.append(
                {
                    "type": "SECTION_THINNER_THAN_REGEN",
                    "section": key,
                    "detail": f"Stored html text much shorter than Python re-split of same content_html "
                    f"(stored ~{len(st_body)} chars vs regen ~{len(rg_body)}). Possible missing body in JSON.",
                    "stored_chars": len(st_body),
                    "regen_chars": len(rg_body),
                }
            )

    keys_with_html = []
    for heading in FIXED_HEADINGS:
        if heading == "Overview":
            continue
        k = section_key(heading)
        block = stored_sections.get(k)
        if isinstance(block, dict):
            t = plain_text(block.get("html") or "")
        elif isinstance(block, str):
            t = plain_text(block)
        else:
            t = ""
        if len(t) > 120:
            keys_with_html.append((k, t))

    for i, (k1, t1) in enumerate(keys_with_html):
        for k2, t2 in keys_with_html[i + 1 :]:
            sim = similarity(t1, t2)
            if sim > 0.88:
                issues.append(
                    {
                        "type": "DUPLICATE_ACROSS_SECTIONS",
                        "sections": [k1, k2],
                        "detail": f"Very similar body text between '{k1}' and '{k2}' (similarity {sim:.2f}).",
                        "similarity": round(sim, 3),
                    }
                )

    pc = stored_sections.get("pros_cons")
    if isinstance(pc, dict):
        for field in ("pros", "cons"):
            inner = pc.get(field)
            if not inner or not str(inner).strip():
                continue
            dups = duplicate_li_titles(inner)
            if dups:
                issues.append(
                    {
                        "type": "DUPLICATE_LIST_ITEMS",
                        "section": "pros_cons",
                        "field": field,
                        "detail": "Repeated bullet lines in %s: %s"
                        % (field, "; ".join(d[:100] for d in dups[:5])),
                    }
                )
        inner_full = pc.get("html")
        if inner_full and str(inner_full).strip():
            soup = BeautifulSoup(inner_full, "html.parser")
            texts = [
                plain_text(str(li))[:200].lower()
                for li in soup.find_all("li")
                if plain_text(str(li))
            ]
            cnt = Counter(texts)
            html_li_dups = [t for t, c in cnt.items() if c > 1 and len(t) > 10]
            if html_li_dups:
                issues.append(
                    {
                        "type": "DUPLICATE_LIST_ITEMS",
                        "section": "pros_cons",
                        "field": "html",
                        "detail": "Repeated <li> text in pros_cons html: " + "; ".join(html_li_dups[:5]),
                    }
                )

    return {"id": cid, "name": name, "issues": issues}


DEFAULT_WORD_ROOT = "/home/itpc6/Public/share/content- Topteen/vocational courses"


class Command(BaseCommand):
    help = (
        "Cross-check vocational course content_html vs content_json display text; "
        "flag missing coverage, thin sections vs regen, cross-section duplicates, duplicate list items. "
        "With --word-root, also compare each course to a matching .docx under that folder (see module docstring)."
    )

    def add_arguments(self, parser):
        parser.add_argument("--course-id", type=int, default=None, help="Only this primary key.")
        parser.add_argument(
            "--output",
            type=str,
            default=None,
            help="Write report to this file (UTF-8). Default: stdout only.",
        )
        parser.add_argument(
            "--format",
            choices=("text", "json"),
            default="text",
            help="Report format (default: text).",
        )
        parser.add_argument(
            "--word-root",
            type=str,
            default=None,
            metavar="DIR",
            help=(
                "Root folder to scan recursively for *.docx. Each file stem is matched to "
                "VocationalCourse.name (normalized). Default when --with-word-docs is used: "
                f"{DEFAULT_WORD_ROOT!r}"
            ),
        )
        parser.add_argument(
            "--with-word-docs",
            action="store_true",
            help=f"Enable Word comparison using --word-root (default directory if not set).",
        )

    def handle(self, *args, **options):
        qs = VocationalCourse.objects.all().order_by("pk")
        if options.get("course_id"):
            qs = qs.filter(pk=options["course_id"])

        word_root_opt = options.get("word_root")
        with_word = options.get("with_word_docs")
        word_root: Path | None = None
        if word_root_opt:
            word_root = Path(word_root_opt).expanduser()
        elif with_word:
            word_root = Path(DEFAULT_WORD_ROOT).expanduser()
        if word_root is not None:
            if not word_root.is_dir():
                self.stderr.write(self.style.ERROR(f"Word root is not a directory: {word_root}"))
                return

        stem_to_path = scan_docx_files(word_root) if word_root else {}
        used_doc_paths: set[str] = set()

        rows = []
        for course in qs:
            row = analyze_course(course)
            if word_root and stem_to_path:
                doc_path, match_note = find_word_doc_path(course.name or "", stem_to_path)
                row["word_file"] = doc_path
                row["word_issues"] = []
                if not doc_path:
                    row["word_issues"].append(
                        {
                            "type": "NO_WORD_DOCUMENT",
                            "detail": f"No .docx found under {word_root} matching course name «{course.name}».",
                        }
                    )
                else:
                    used_doc_paths.add(doc_path)
                    try:
                        doc_plain = extract_docx_plain(doc_path)
                    except Exception as e:
                        row["word_issues"].append(
                            {
                                "type": "WORD_READ_ERROR",
                                "detail": str(e),
                                "path": doc_path,
                            }
                        )
                    else:
                        row["word_issues"].extend(
                            word_doc_issues_for_course(course, doc_path, doc_plain, match_note)
                        )
            else:
                row["word_file"] = None
                row["word_issues"] = []
            rows.append(row)

        orphan_docs: list[str] = []
        # Only meaningful when every course was scanned (otherwise "orphan" = every other file).
        if word_root and stem_to_path and not options.get("course_id"):
            all_paths = set(stem_to_path.values())
            orphan_docs = sorted(all_paths - used_doc_paths)

        fmt = options["format"]
        out_lines = []
        if fmt == "json":
            payload = {
                "courses": rows,
                "orphan_word_files": orphan_docs,
                "word_root": str(word_root) if word_root else None,
            }
            out_lines.append(json.dumps(payload, indent=2, ensure_ascii=False))
        else:
            out_lines.append("Vocational course display cross-check")
            if word_root:
                out_lines.append(f"Word .docx root: {word_root}")
            out_lines.append("=" * 72)
            total_issues = 0
            total_word_issues = 0
            for r in rows:
                out_lines.append(f"\n[{r['id']}] {r['name']}")
                if not r["issues"]:
                    out_lines.append("  (DB) OK — no flags.")
                else:
                    for it in r["issues"]:
                        total_issues += 1
                        t = it.get("type", "?")
                        d = it.get("detail", "")
                        out_lines.append(f"  • [{t}] {d}")
                if word_root:
                    wf = r.get("word_file")
                    if wf:
                        out_lines.append(f"  Word file: {wf}")
                    wiss = r.get("word_issues") or []
                    if not wiss:
                        out_lines.append("  (Word) OK — no flags.")
                    else:
                        for it in wiss:
                            total_word_issues += 1
                            t = it.get("type", "?")
                            d = it.get("detail", "")
                            out_lines.append(f"  • [{t}] {d}")
            if orphan_docs:
                out_lines.append("\n--- Orphan .docx (no matching VocationalCourse.name; run without --course-id) ---")
                for op in orphan_docs:
                    out_lines.append(f"  • {op}")
            elif word_root and options.get("course_id"):
                out_lines.append("\n(Orphan .docx list omitted when using --course-id.)")
            out_lines.append("\n" + "=" * 72)
            out_lines.append(
                f"Scanned {len(rows)} course(s); "
                f"{sum(len(r['issues']) for r in rows)} DB issue row(s); "
                f"{sum(len(r.get('word_issues') or []) for r in rows)} Word issue row(s)."
            )
            if orphan_docs:
                out_lines.append(f"Orphan Word files: {len(orphan_docs)}.")

        text = "\n".join(out_lines)
        self.stdout.write(text)
        path = options.get("output")
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write(text)
            self.stdout.write(self.style.SUCCESS(f"\nWrote report to {path}"))
