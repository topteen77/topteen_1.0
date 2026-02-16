"""
One-off script: extract assessment questions from docx files into JSON.
Run from project root: python -m core.four_pillars_assessments.extract_docx_to_json

Captures: questions (title, text, options), profiles (name, summary, scoring_bullets from Scoring Guide),
and preserves bold headings when extracting the Scoring Guide section.

Interest Drivers: scoring guide is read from reference HTML when available:
  topteenhtml/html/a/interest-drivers-assessement.html (sibling of project dir).
"""
import html
import json
import re
from pathlib import Path

# #region agent log
DEBUG_LOG_PATH = Path(__file__).resolve().parent.parent.parent / ".cursor" / "debug.log"
def _debug_log(msg, data, hypothesis_id="H1"):
    try:
        with open(DEBUG_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps({"message": msg, "data": data, "hypothesisId": hypothesis_id, "timestamp": __import__("time").time() * 1000}) + "\n")
    except Exception:
        pass
# #endregion

try:
    from docx import Document
except ImportError:
    print("Install python-docx: pip install python-docx")
    raise

ASSESSMENTS_DIR = Path("/home/itpc6/Public/share/content- Topteen/The Four Pillars Learning Framework/assessments")
OUT_DIR = Path(__file__).resolve().parent

# Reference HTML for Interest Drivers scoring guide (single source of truth)
# Path: 7nov/topteenhtml/html/a/interest-drivers-assessement.html (7nov = 5 levels up from this file)
_SCRIPT_DIR = Path(__file__).resolve().parent
_REPO_ROOT = _SCRIPT_DIR.parent.parent.parent.parent.parent  # 7nov
INTEREST_DRIVERS_REFERENCE_HTML = _REPO_ROOT / "topteenhtml" / "html" / "a" / "interest-drivers-assessement.html"
ENGAGEMENT_PATTERNS_REFERENCE_HTML = _REPO_ROOT / "topteenhtml" / "html" / "a" / "engagement-patterns.html"
LEARNING_PREFERENCES_REFERENCE_HTML = _REPO_ROOT / "topteenhtml" / "html" / "a" / "learning-preferences-assessment.html"
NATURAL_ABILITIES_REFERENCE_HTML = _REPO_ROOT / "topteenhtml" / "html" / "a" / "natural-abilities -assessment.html"  # note space in filename
# Where to write Engagement Patterns "Understanding" static files (project dir = topteen_1.0)
_PROJECT_ROOT = _SCRIPT_DIR.parent.parent
ENGAGEMENT_UNDERSTANDING_INCLUDE = _PROJECT_ROOT / "templates" / "template20" / "includes" / "engagement_patterns_understanding_guide.html"
ENGAGEMENT_UNDERSTANDING_JSON = _SCRIPT_DIR / "engagement_patterns_understanding_data.json"
# Config for all four "understanding data" JSON files: (slug, html_path, title, json_filename)
UNDERSTANDING_DATA_CONFIG = [
    ("learning_preferences", LEARNING_PREFERENCES_REFERENCE_HTML, "Understanding Your Learning Profile", "learning_preferences_understanding_data.json"),
    ("natural_abilities", NATURAL_ABILITIES_REFERENCE_HTML, "Understanding Your Natural Abilities", "natural_abilities_understanding_data.json"),
    ("engagement_patterns", ENGAGEMENT_PATTERNS_REFERENCE_HTML, "Understanding Your Engagement Patterns", "engagement_patterns_understanding_data.json"),
    ("interest_drivers", INTEREST_DRIVERS_REFERENCE_HTML, "Understanding Your Interest Drivers", "interest_drivers_understanding_data.json"),
]

FILES = [
    ("Learning Preferences Assessment questions.docx", "learning_preferences", "Learning Preferences"),
    ("Natural Abilities Assessment Questions.docx", "natural_abilities", "Natural Abilities"),
    ("Engagement Patterns assessment.docx", "engagement_patterns", "Engagement Patterns"),
    ("Interest Drivers assessment.docx", "interest_drivers", "Interest Drivers"),
]

# Default profile labels per pillar (A/B/C/D)
DEFAULT_PROFILES = {
    "learning_preferences": {
        "A": {"name": "Scholarly Learner (Theoretical/Reading-focused)", "summary": "You learn best through deep reading, research, and careful analysis. Quiet study spaces, detailed notes, and comprehensive understanding are your strengths."},
        "B": {"name": "Experiential Learner (Kinesthetic/Hands-on)", "summary": "You learn best by doing. Hands-on activities, real-world applications, and experimentation help you understand and remember information."},
        "C": {"name": "Social Learner (Auditory/Collaborative)", "summary": "You learn best through discussion, collaboration, and verbal processing. Study groups, conversations, and teaching others help ideas stick."},
        "D": {"name": "Structured Learner (Visual/Organized)", "summary": "You learn best with clear structure, visual organization, and step-by-step guidance. Plans, checklists, and visual tools support your success."},
    },
    "natural_abilities": {
        "A": {"name": "Analytical & Technical", "summary": "You excel at pattern recognition, logical analysis, and systematic problem-solving. You naturally break down complexity and find structure."},
        "B": {"name": "Practical & Applied", "summary": "You thrive when completing concrete tasks and seeing tangible results. You are motivated by getting things done and real-world impact."},
        "C": {"name": "Communication & Interpersonal", "summary": "You naturally connect with others, ask probing questions, and communicate ideas clearly. You learn and contribute through dialogue."},
        "D": {"name": "Creative & Innovative", "summary": "You are drawn to new ideas, creative solutions, and exploring possibilities. You think in terms of what could be rather than only what is."},
    },
    "engagement_patterns": {
        "A": {"name": "Goal-Oriented Achiever", "summary": "You are energized by clear goals, measurable progress, and achieving specific outcomes. You like steady progress and milestones."},
        "B": {"name": "Hands-On Creator", "summary": "You are motivated by building, creating, and making things. You prefer learning by doing and seeing tangible results."},
        "C": {"name": "Knowledge Seeker", "summary": "You are driven by understanding concepts deeply and exploring ideas. You value insight and mastery of subject matter."},
        "D": {"name": "Balanced Innovator", "summary": "You combine theory and practice. You like both understanding why and applying how, with flexibility in your approach."},
    },
    "interest_drivers": {
        "A": {"name": "Analytical & Research-Focused", "summary": "You are drawn to data, facts, research, and logical analysis. You want to understand evidence and how conclusions are reached."},
        "B": {"name": "Practical & Systems-Focused", "summary": "You are interested in how things work, technical details, and practical applications. You like solving real problems."},
        "C": {"name": "People & Experience-Focused", "summary": "You are curious about people, stories, and human experience. You engage with ideas through conversation and narrative."},
        "D": {"name": "Creative & Big-Picture Focused", "summary": "You are drawn to creative expression, trends, and broader implications. You like connecting ideas and seeing the bigger picture."},
    },
}

# Canonical short Scoring Guide for Learning Preferences (displayed on the Scoring Guide section only)
CANONICAL_SCORING_GUIDE_LEARNING_PREFERENCES = {
    "intro": "Scoring Guide: Count your responses for each letter:",
    "A": {
        "heading": "Mostly A's: Scholarly Learner (Theoretical/Reading-focused)",
        "items": [
            "Prefers deep reading and research",
            "Thrives with written materials and quiet study",
            "Values thorough understanding and academic rigor",
            "Learns best through comprehensive analysis and reflection",
        ],
    },
    "B": {
        "heading": "Mostly B's: Experiential Learner (Kinesthetic/Hands-on)",
        "items": [
            "Learns best through hands-on experience and experimentation",
            "Prefers active, practical applications over theory",
            "Motivated by real-world connections and immediate application",
            "Thrives in dynamic, interactive learning environments",
        ],
    },
    "C": {
        "heading": "Mostly C's: Social Learner (Auditory/Collaborative)",
        "items": [
            "Learns through discussion, collaboration, and verbal processing",
            "Values different perspectives and group interactions",
            "Motivated by interpersonal connections and shared learning",
            "Processes information best through talking and listening",
        ],
    },
    "D": {
        "heading": "Mostly D's: Structured Learner (Visual/Organized)",
        "items": [
            "Prefers organized, systematic approaches to learning",
            "Values clear guidance, visual aids, and measurable progress",
            "Thrives with balanced, methodical learning processes",
            "Learns best with clear structure and visual organization",
        ],
    },
    "mixed_results": "Mixed Results: Many learners have a combination of preferences. Look at your top two categories to understand your primary and secondary learning styles.",
}

# Canonical short Scoring Guide for Engagement Patterns (Scoring Guide section on assessment page)
CANONICAL_SCORING_GUIDE_ENGAGEMENT_PATTERNS = {
    "intro": (
        "Scoring Guide: How to Calculate Your Score: "
        "Step 1: Count your responses for each letter: "
        "Count all your A responses (Achievement-Driven), B (Mastery-Oriented), C (Purpose-Driven), D (Variety-Seeking). "
        "Step 2: Determine your engagement pattern based on your highest scores."
    ),
    "A": {
        "heading": "A) Achievement-Driven (Results & Goal Orientation)",
        "items": [
            "Goal Completion: Deeply motivated by reaching specific targets and objectives",
            "Progress Visibility: Need to see measurable advancement and concrete results",
            "Performance Excellence: Driven by the satisfaction of high performance and success",
            "Competitive Edge: Energized by challenges that test and prove capabilities",
        ],
    },
    "B": {
        "heading": "B) Mastery-Oriented (Learning & Expertise Development)",
        "items": [
            "Skill Development: Deeply motivated by becoming excellent at meaningful work",
            "Knowledge Acquisition: Energized by learning new concepts and expanding understanding",
            "Expertise Building: Driven by the journey toward mastery and professional growth",
            "Intellectual Challenge: Motivated by complex problems that require deep thinking",
        ],
    },
    "C": {
        "heading": "C) Purpose-Driven (Meaning & Impact Orientation)",
        "items": [
            "Meaningful Impact: Deeply motivated by work that makes a positive difference",
            "Value Alignment: Energized when work connects to personal values and beliefs",
            "Service to Others: Driven by opportunities to help, support, and contribute to others",
            "Legacy Building: Motivated by creating lasting positive change",
        ],
    },
    "D": {
        "heading": "D) Variety-Seeking (Stimulation & Change Orientation)",
        "items": [
            "Novelty and Change: Deeply motivated by new experiences and fresh challenges",
            "Stimulation and Excitement: Energized by dynamic, unpredictable environments",
            "Flexibility and Autonomy: Driven by freedom to adapt and try different approaches",
            "Growth Through Challenge: Motivated by overcoming obstacles and expanding comfort zones",
        ],
    },
    "mixed_results": (
        "Dual Engagement Pattern (Two Dominant Styles): Score Range 8–12 in two categories, others below 6. "
        "e.g. High A + High B: Achievement-Mastery (Expert Achiever); High A + High C: Achievement-Purpose (Mission-Focused Achiever); "
        "High B + High C: Mastery-Purpose (Meaningful Expert); High C + High D: Purpose-Variety (Flexible Contributor). "
        "Balanced Engagement Pattern: 6–10 across three or more categories. "
        "Multi-Modal Engagement: Relatively even scores across all categories (5–8 in each)."
    ),
}

# Canonical short Scoring Guide for Interest Drivers (from topteenhtml/html/a/interest-drivers-assessement.html)
CANONICAL_SCORING_GUIDE_INTEREST_DRIVERS = {
    "intro": (
        "Step 1: Count your responses for each letter: A (Analytical Interest), B (Technical Interest), C (People Interest), D (Creative Interest). "
        "Step 2: Determine your interest driver pattern from the ranges below."
    ),
    "A": {
        "heading": "15–20 A's: Strong Analytical Interest",
        "items": [
            "Data & research curiosity. Evidence, proof, logical investigation; thrives with research databases, systematic analysis, and evidence-based content.",
        ],
    },
    "B": {
        "heading": "15–20 B's: Strong Technical Interest",
        "items": [
            "Systems & process curiosity. How things work, practical application, skill development; thrives with hands-on learning and technical problem-solving.",
        ],
    },
    "C": {
        "heading": "15–20 C's: Strong People Interest",
        "items": [
            "Human & social curiosity. Human behavior, relationship dynamics, social impact; thrives with stories, collaboration, and understanding people.",
        ],
    },
    "D": {
        "heading": "15–20 D's: Strong Creative Interest",
        "items": [
            "Innovation & possibility curiosity. New ideas, creative solutions, future orientation; thrives with creative expression and exploring possibilities.",
        ],
    },
    "mixed_results": (
        "Dual Interest Driver: 8–12 in two categories (e.g. A+B Research Engineer, C+D Social Innovator). "
        "Balanced: 6–10 across three or more = combination pattern. "
        "Multi-Domain Curiosity: 5–8 in each = balanced across all four. See combination profiles below."
    ),
}


def _strip_html(text: str) -> str:
    """Remove HTML tags and decode entities."""
    if not text:
        return ""
    text = re.sub(r"<[^>]+>", " ", text)
    text = " ".join(text.split())
    return html.unescape(text).strip()


def extract_interest_drivers_scoring_from_html(html_path: Path) -> dict | None:
    """
    Parse the Interest Drivers scoring guide from the reference HTML.
    Returns same structure as CANONICAL_SCORING_GUIDE_INTEREST_DRIVERS, or None on failure.
    """
    if not html_path.exists():
        return None
    try:
        raw = html_path.read_text(encoding="utf-8")
    except Exception:
        return None
    # Intro: <p><strong>Step 1:</strong> ... <strong>Step 2:</strong> ... </p>
    intro_m = re.search(
        r"<p><strong>Step\s+1:</strong>\s*([^<]+)\.\s*<strong>Step\s+2:</strong>\s*([^<]+)\.?</p>",
        raw,
        re.DOTALL,
    )
    if not intro_m:
        return None
    intro = "Step 1: " + html.unescape(intro_m.group(1).strip()) + ". Step 2: " + html.unescape(intro_m.group(2).strip()) + "."

    # Cards: .lp-scoring-card a/b/c/d with .lp-scoring-card-title and .lp-scoring-card-desc
    cards_section = re.search(
        r'<div class="lp-scoring-cards">(.*?)</div>\s*<div class="lp-mixed-note">',
        raw,
        re.DOTALL,
    )
    if not cards_section:
        return None
    block = cards_section.group(1)
    scoring = {
        "intro": intro,
        "A": {"heading": "", "items": []},
        "B": {"heading": "", "items": []},
        "C": {"heading": "", "items": []},
        "D": {"heading": "", "items": []},
        "mixed_results": "",
    }
    for letter in ("A", "B", "C", "D"):
        card = re.search(
            rf'<div class="lp-scoring-card {letter.lower()}">.*?'
            r'<div class="lp-scoring-card-title">([^<]+)</div>'
            r'.*?<div class="lp-scoring-card-desc">([^<]*)</div>',
            block,
            re.DOTALL,
        )
        if card:
            scoring[letter]["heading"] = html.unescape(card.group(1).strip())
            desc = html.unescape(card.group(2).strip())
            if desc:
                scoring[letter]["items"] = [desc]

    # Mixed note: <div class="lp-mixed-note">...</div>
    mixed_m = re.search(r'<div class="lp-mixed-note">(.*?)</div>', raw, re.DOTALL)
    if mixed_m:
        scoring["mixed_results"] = _strip_html(mixed_m.group(1))

    if not all(scoring[k]["heading"] and scoring[k]["items"] for k in ("A", "B", "C", "D")):
        return None
    if not scoring["mixed_results"]:
        return None
    return scoring


def _extract_understanding_fragment(raw: str, h3_end_marker: str) -> str | None:
    """Extract the 'Understanding Your X' block (h3 + paragraph + accordion). h3_end_marker e.g. 'Understanding Your Engagement Patterns</h3>'."""
    start = raw.find(h3_end_marker)
    if start == -1:
        return None
    start = raw.rfind("<h3", 0, start)
    if start == -1:
        return None
    acc_start = raw.find('<div class="accordion lp-guide-accordion" id="lpGuideAccordion">', start)
    if acc_start == -1:
        return None
    depth = 0
    i = acc_start
    while i < len(raw):
        if raw[i : i + 5] == "<div " or raw[i : i + 5] == "<div>":
            depth += 1
            i += 4
            continue
        if raw[i : i + 6] == "</div>":
            depth -= 1
            if depth == 0:
                return raw[start : i + 6]
            i += 6
            continue
        i += 1
    return None


def _extract_engagement_understanding_fragment(raw: str) -> str | None:
    """Extract the 'Understanding Your Engagement Patterns' block (h3 + paragraph + accordion)."""
    return _extract_understanding_fragment(raw, "Understanding Your Engagement Patterns</h3>")


def _extract_primary_section_lists(html_fragment: str) -> list:
    """From lp-primary-box content, extract sections with heading + bullet list (same structure as HTML)."""
    sections = []
    # Each lp-primary-section: <div class="lp-primary-section"> <strong>...</strong> <ul> <li>...</li> ... </ul> </div>
    pat = re.compile(
        r'<div class="lp-primary-section">\s*<strong>([^<]*)</strong>\s*<ul>\s*(.*?)\s*</ul>\s*</div>',
        re.DOTALL,
    )
    for m in pat.finditer(html_fragment):
        heading = html.unescape(m.group(1).strip())
        ul_content = m.group(2)
        items = re.findall(r"<li>(.*?)</li>", ul_content, re.DOTALL)
        # Keep bold/italic/list format in each item (same as HTML)
        items = [html.unescape(x.strip()) for x in items]
        sections.append({"heading": heading, "bullets": items})
    return sections


def generate_engagement_understanding_static_files() -> bool:
    """
    Read engagement-patterns.html and generate:
    1) Static HTML fragment for 'Understanding Your Engagement Patterns' (bold, italic, lists preserved).
    2) JSON data file with same bullet list structure for use in frontend/data.
    """
    if not ENGAGEMENT_PATTERNS_REFERENCE_HTML.exists():
        print("Skip engagement understanding (HTML not found):", ENGAGEMENT_PATTERNS_REFERENCE_HTML)
        return False
    raw = ENGAGEMENT_PATTERNS_REFERENCE_HTML.read_text(encoding="utf-8")
    fragment = _extract_engagement_understanding_fragment(raw)
    if not fragment:
        print("Skip engagement understanding (fragment not found in HTML)")
        return False
    ENGAGEMENT_UNDERSTANDING_INCLUDE.parent.mkdir(parents=True, exist_ok=True)
    ENGAGEMENT_UNDERSTANDING_INCLUDE.write_text(fragment, encoding="utf-8")
    print("Wrote", ENGAGEMENT_UNDERSTANDING_INCLUDE)

    # Build JSON with same bullet list structure (intro + primary profiles with sections as bullet lists)
    intro_p = re.search(
        r"<p>Your engagement pattern reveals[^<]+</p>",
        fragment,
    )
    intro = html.unescape(intro_p.group(0).replace("<p>", "").replace("</p>", "").strip()) if intro_p else ""

    data = {"title": "Understanding Your Engagement Patterns", "intro": intro, "primary_profiles": []}
    # Extract each lp-primary-box (A, B, C, D)
    for letter in ("a", "b", "c", "d"):
        box = re.search(
            rf'<div class="lp-primary-box {letter}">(.*?)</div>\s*</div>\s*</div>\s*</div>\s*</div>',
            fragment,
            re.DOTALL,
        )
        if not box:
            continue
        box_html = box.group(1)
        h4 = re.search(r"<h4>([^<]+)</h4>", box_html)
        label = html.unescape(h4.group(1).strip()) if h4 else f"Profile {letter.upper()}"
        sections = _extract_primary_section_lists(box_html)
        data["primary_profiles"].append({"label": label, "sections": sections})

    with open(ENGAGEMENT_UNDERSTANDING_JSON, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print("Wrote", ENGAGEMENT_UNDERSTANDING_JSON)
    return True


def _extract_lp_primary_box_html(fragment: str, letter: str) -> str | None:
    """Extract inner HTML of first <div class='lp-primary-box {letter}'> by matching div depth."""
    pattern = rf'<div class="lp-primary-box {letter}">'
    start = fragment.find(pattern)
    if start == -1:
        return None
    start = fragment.find(">", start) + 1
    depth = 1
    i = start
    while i < len(fragment):
        if fragment[i : i + 5] == "<div " or fragment[i : i + 5] == "<div>":
            depth += 1
            i += 4
            continue
        if fragment[i : i + 6] == "</div>":
            depth -= 1
            if depth == 0:
                return fragment[start:i]
            i += 6
            continue
        i += 1
    return None


def _build_understanding_data_from_fragment(fragment: str, title: str) -> dict:
    """Build understanding_data JSON dict from HTML fragment (intro + primary_profiles with sections/bullets)."""
    intro_p = re.search(r"<p>(.*?)</p>", fragment, re.DOTALL)
    intro = html.unescape(re.sub(r"<[^>]+>", " ", (intro_p.group(1) or "").strip())).strip() if intro_p else ""
    data = {"title": title, "intro": intro, "primary_profiles": []}
    for letter in ("a", "b", "c", "d"):
        box_html = _extract_lp_primary_box_html(fragment, letter)
        if not box_html:
            continue
        h4 = re.search(r"<h4>([^<]+)</h4>", box_html)
        label = html.unescape(h4.group(1).strip()) if h4 else f"Profile {letter.upper()}"
        sections = _extract_primary_section_lists(box_html)
        data["primary_profiles"].append({"label": label, "sections": sections})
    return data


def generate_all_understanding_data():
    """
    Generate understanding_data.json for all four assessments from reference HTML.
    Reads each HTML, extracts the 'Understanding Your X' block, builds JSON with same
    bullet list structure, writes to core/four_pillars_assessments/<slug>_understanding_data.json.
    For engagement_patterns also writes the static HTML include.
    """
    h3_markers = {
        "learning_preferences": "Understanding Your Learning Profile</h3>",
        "natural_abilities": "Understanding Your Natural Abilities</h3>",
        "engagement_patterns": "Understanding Your Engagement Patterns</h3>",
        "interest_drivers": "Understanding Your Interest Drivers</h3>",
    }
    for _slug, html_path, title, json_filename in UNDERSTANDING_DATA_CONFIG:
        if not html_path.exists():
            print("Skip", json_filename, "(HTML not found):", html_path)
            continue
        raw = html_path.read_text(encoding="utf-8")
        marker = h3_markers.get(_slug, title.replace("Your ", "").replace(" ", "") + "</h3>")
        if marker.count("</h3>") != 1:
            marker = title + "</h3>"
        fragment = _extract_understanding_fragment(raw, marker)
        if not fragment:
            print("Skip", json_filename, "(fragment not found)")
            continue
        data = _build_understanding_data_from_fragment(fragment, title)
        out_path = _SCRIPT_DIR / json_filename
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print("Wrote", out_path)
        if _slug == "engagement_patterns":
            ENGAGEMENT_UNDERSTANDING_INCLUDE.parent.mkdir(parents=True, exist_ok=True)
            ENGAGEMENT_UNDERSTANDING_INCLUDE.write_text(fragment, encoding="utf-8")
            print("Wrote", ENGAGEMENT_UNDERSTANDING_INCLUDE)


def paragraph_text(para) -> str:
    """Full text of paragraph."""
    return (para.text or "").strip()


def paragraph_is_bold(para) -> bool:
    """True if paragraph has at least one run and all non-empty runs are bold."""
    if not para.runs:
        return False
    for run in para.runs:
        if (run.text or "").strip():
            if not getattr(run, "bold", False):
                return False
    return True


def is_list_paragraph(para) -> bool:
    """Heuristic: paragraph is a list item if style name suggests it or numPr present."""
    try:
        name = (para.style and para.style.name) or ""
        if name and ("List" in name or "Bullet" in name):
            return True
    except Exception:
        pass
    try:
        pPr = para._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        if pPr is not None and pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr") is not None:
            return True
    except Exception:
        pass
    return False


# Patterns for Scoring Guide section in DOCX (all four assessment types)
# Learning Preferences: "Mostly A's:", "Mixed Results:"
# Natural Abilities / Engagement / Interest: "A) Profile Name", "Step 1: Count", "Dual Profile", etc.
RE_MOSTLY_A = re.compile(r"^\s*Mostly\s+A'?s\s*:?", re.I)
RE_MOSTLY_B = re.compile(r"^\s*Mostly\s+B'?s\s*:?", re.I)
RE_MOSTLY_C = re.compile(r"^\s*Mostly\s+C'?s\s*:?", re.I)
RE_MOSTLY_D = re.compile(r"^\s*Mostly\s+D'?s\s*:?", re.I)
# Generic profile heading: "A) Name", "A. Name", "Category A:", "8–15 in A:", "Mostly A's:" etc.
RE_HEADING_A = re.compile(r"^\s*(?:Mostly\s+A'?s\s*:?|A\s*[).]|Category\s+A\s*:?|Profile\s+A\s*:?|\d+[–\-]\d+\s+in\s+A\s*:?)", re.I)
RE_HEADING_B = re.compile(r"^\s*(?:Mostly\s+B'?s\s*:?|B\s*[).]|Category\s+B\s*:?|Profile\s+B\s*:?|\d+[–\-]\d+\s+in\s+B\s*:?)", re.I)
RE_HEADING_C = re.compile(r"^\s*(?:Mostly\s+C'?s\s*:?|C\s*[).]|Category\s+C\s*:?|Profile\s+C\s*:?|\d+[–\-]\d+\s+in\s+C\s*:?)", re.I)
RE_HEADING_D = re.compile(r"^\s*(?:Mostly\s+D'?s\s*:?|D\s*[).]|Category\s+D\s*:?|Profile\s+D\s*:?|\d+[–\-]\d+\s+in\s+D\s*:?)", re.I)
# End of scoring section (mixed/dual/balanced text - not "Combination Learning Profiles" which is a section heading)
RE_MIXED = re.compile(
    r"^\s*(?:Mixed\s+Results\s*:?|Dual\s+(?:Profile|Engagement|Interest)\s*:?|Balanced\s*:?|Multi[- ]?(?:Modal|Domain)\s*[:(]?)",
    re.I,
)

# Section intro triggers (any of these start the scoring guide)
SCORING_INTRO_TRIGGERS = ("Scoring Guide", "Count your responses", "Step 1: Count", "Step 2:")


def extract_scoring_guide(doc) -> dict:
    """
    Extract Scoring Guide section: intro, for each of A/B/C/D a heading and list of bullet items, and mixed/dual results text.
    Supports Learning Preferences (Mostly A's), Natural Abilities, Engagement, Interest (A), B), etc.).
    Returns dict with intro, A/B/C/D each { "heading": str, "items": list }, mixed_results: str.
    """
    result = {
        "intro": "",
        "A": {"heading": "", "items": []},
        "B": {"heading": "", "items": []},
        "C": {"heading": "", "items": []},
        "D": {"heading": "", "items": []},
        "mixed_results": "",
    }
    paras = list(doc.paragraphs)
    i = 0
    # Find start: any known scoring-section intro
    while i < len(paras):
        t = paragraph_text(paras[i])
        if any(trigger in t for trigger in SCORING_INTRO_TRIGGERS):
            parts = [t]
            # Optionally include next paragraph if it continues the intro (e.g. "Step 2: ...")
            if i + 1 < len(paras):
                next_t = paragraph_text(paras[i + 1])
                if next_t and any(trigger in next_t for trigger in ("Step 2:", "Count your responses", "Use the ranges")):
                    parts.append(next_t)
                    i += 1
            result["intro"] = " ".join(parts)
            i += 1
            break
        i += 1
    if i >= len(paras):
        return result
    # Scan for profile headings A/B/C/D (multiple patterns) and mixed/dual section
    current = None  # "A" | "B" | "C" | "D" | "mixed"
    prev_current = None
    j = i
    while j < len(paras):
        p = paras[j]
        t = paragraph_text(p)
        if not t:
            j += 1
            continue
        if RE_HEADING_A.search(t):
            # #region agent log
            if current:
                _debug_log("switch_heading", {"from": current, "to": "A", "items_count": len(result.get(current, {}).get("items", []))}, "H3")
            _debug_log("start_profile", {"current": "A", "j": j, "heading_preview": t[:60], "A_already_has_items": len(result["A"]["items"])}, "H1")
            # #endregion
            prev_current = current
            if len(result["A"]["items"]) == 0:
                current = "A"
                result["A"]["heading"] = t
            else:
                current = None
            j += 1
            continue
        if RE_HEADING_B.search(t):
            # #region agent log
            if current:
                _debug_log("switch_heading", {"from": current, "to": "B", "items_count": len(result.get(current, {}).get("items", []))}, "H3")
            _debug_log("start_profile", {"current": "B", "j": j, "heading_preview": t[:60], "B_already_has_items": len(result["B"]["items"])}, "H1")
            # #endregion
            prev_current = current
            if len(result["B"]["items"]) == 0:
                current = "B"
                result["B"]["heading"] = t
            else:
                current = None
            j += 1
            continue
        if RE_HEADING_C.search(t):
            # #region agent log
            if current:
                _debug_log("switch_heading", {"from": current, "to": "C", "items_count": len(result.get(current, {}).get("items", []))}, "H3")
            _debug_log("start_profile", {"current": "C", "j": j, "heading_preview": t[:60], "C_already_has_items": len(result["C"]["items"])}, "H1")
            # #endregion
            prev_current = current
            if len(result["C"]["items"]) == 0:
                current = "C"
                result["C"]["heading"] = t
            else:
                current = None
            j += 1
            continue
        if RE_HEADING_D.search(t):
            # #region agent log
            if current:
                _debug_log("switch_heading", {"from": current, "to": "D", "items_count": len(result.get(current, {}).get("items", []))}, "H3")
            _debug_log("start_profile", {"current": "D", "j": j, "heading_preview": t[:60], "D_already_has_items": len(result["D"]["items"])}, "H1")
            # #endregion
            prev_current = current
            if len(result["D"]["items"]) == 0:
                current = "D"
                result["D"]["heading"] = t
            else:
                current = None
            j += 1
            continue
        if RE_MIXED.search(t):
            # #region agent log
            if current:
                _debug_log("switch_heading", {"from": current, "to": "mixed", "items_count": len(result.get(current, {}).get("items", []))}, "H3")
            # #endregion
            result["mixed_results"] = t
            # Allow only the immediately next paragraph (short mixed message)
            if j + 1 < len(paras):
                next_t = paragraph_text(paras[j + 1])
                if next_t and not RE_HEADING_A.search(next_t) and not RE_HEADING_B.search(next_t) and not RE_HEADING_C.search(next_t) and not RE_HEADING_D.search(next_t) and not RE_MIXED.search(next_t):
                    result["mixed_results"] += " " + next_t
                    j += 1
            current = None
            j += 1
            continue
        if current in ("A", "B", "C", "D"):
            # #region agent log
            is_list = is_list_paragraph(p)
            _debug_log("append_item", {"j": j, "current": current, "is_list": is_list, "text_preview": t[:55], "items_so_far": len(result[current]["items"]) + (1 if is_list else 0)}, "H2")
            # #endregion
            if is_list_paragraph(p):
                result[current]["items"].append(t)
        j += 1
    # #region agent log
    _debug_log("scoring_totals", {k: len(result.get(k, {}).get("items", [])) for k in ("A", "B", "C", "D")}, "H5")
    # #endregion
    return result


def parse_options_line(line: str) -> dict:
    """Parse 'A) ... B) ... C) ... D) ...' into { A: ..., B: ..., C: ..., D: ... }."""
    line = line.replace("\n", " ").strip()
    # Split by " B) ", " C) ", " D) " keeping the delimiter with next part
    parts = re.split(r"\s+(?=B\))\s*", line, maxsplit=1)
    a_text = parts[0].strip()
    if a_text.startswith("A)"):
        a_text = a_text[2:].strip()
    rest = parts[1] if len(parts) > 1 else ""
    parts = re.split(r"\s+(?=C\))\s*", rest, maxsplit=1)
    b_text = parts[0].strip()
    if b_text.startswith("B)"):
        b_text = b_text[2:].strip()
    rest = parts[1] if len(parts) > 1 else ""
    parts = re.split(r"\s+(?=D\))\s*", rest, maxsplit=1)
    c_text = parts[0].strip()
    if c_text.startswith("C)"):
        c_text = c_text[2:].strip()
    d_text = parts[1].strip() if len(parts) > 1 else ""
    if d_text.startswith("D)"):
        d_text = d_text[2:].strip()
    return {"A": a_text, "B": b_text, "C": c_text, "D": d_text}


def extract_questions(doc) -> list:
    """Extract questions from document. `doc` is a python-docx Document (or path)."""
    if isinstance(doc, (Path, str)):
        doc = Document(doc)
    paras = [paragraph_text(p) for p in doc.paragraphs if paragraph_text(p)]
    # Skip main title (and optional subtitle)
    start = 0
    if paras and ("Question 1" not in paras[0]):
        start = 1
    if start < len(paras) and "Question 1" not in paras[start]:
        start += 1
    questions = []
    i = start
    while i + 2 < len(paras):
        title = paras[i]
        text = paras[i + 1]
        options_line = paras[i + 2]
        if not title.startswith("Question ") or "A)" not in options_line:
            i += 1
            continue
        try:
            options = parse_options_line(options_line)
        except Exception:
            i += 1
            continue
        questions.append({"title": title, "text": text, "options": options})
        i += 3
    return questions


def merge_scoring_into_profiles(profiles: dict, scoring: dict) -> None:
    """In-place: add scoring_heading and scoring_bullets to each profile A/B/C/D; add scoring_intro and mixed_results to data later."""
    for key in ("A", "B", "C", "D"):
        if key not in profiles:
            profiles[key] = {}
        if scoring.get(key, {}).get("heading"):
            profiles[key]["scoring_heading"] = scoring[key]["heading"]
        items = scoring.get(key, {}).get("items") or []
        if items:
            profiles[key]["scoring_bullets"] = items


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if not ASSESSMENTS_DIR.exists():
        print("Note: Assessments dir not found:", ASSESSMENTS_DIR)
        print("Will generate/update all four JSONs using existing JSON + canonical scoring where DOCX is missing.\n")

    for filename, slug, _ in FILES:
        path = ASSESSMENTS_DIR / filename
        existing_path = OUT_DIR / f"{slug}.json"
        data_existing = None
        if path.exists():
            doc = Document(path)
            questions = extract_questions(doc)
            scoring = extract_scoring_guide(doc)
            # #region agent log
            _debug_log("after_extract_per_file", {"slug": slug, "A_items": len(scoring.get("A", {}).get("items", [])), "B_items": len(scoring.get("B", {}).get("items", [])), "C_items": len(scoring.get("C", {}).get("items", [])), "D_items": len(scoring.get("D", {}).get("items", []))}, "H5")
            # #endregion
        else:
            # No DOCX: keep questions and profiles from existing JSON when present
            if existing_path.exists():
                with open(existing_path, "r", encoding="utf-8") as f:
                    data_existing = json.load(f)
                questions = data_existing.get("questions", [])
            else:
                questions = []
            scoring = {}

        # Apply canonical scoring for learning_preferences, engagement_patterns, interest_drivers
        if slug == "learning_preferences":
            scoring = dict(CANONICAL_SCORING_GUIDE_LEARNING_PREFERENCES)
        elif slug == "engagement_patterns":
            scoring = dict(CANONICAL_SCORING_GUIDE_ENGAGEMENT_PATTERNS)
        elif slug == "interest_drivers":
            from_html = extract_interest_drivers_scoring_from_html(INTEREST_DRIVERS_REFERENCE_HTML)
            if from_html:
                scoring = from_html
            else:
                scoring = dict(CANONICAL_SCORING_GUIDE_INTEREST_DRIVERS)

        base = DEFAULT_PROFILES.get(slug, DEFAULT_PROFILES["learning_preferences"])
        # When DOCX missing and we loaded existing JSON, keep its profiles; else use defaults
        if data_existing is not None and data_existing.get("profiles"):
            profiles = {k: dict(v) for k, v in data_existing["profiles"].items()}
            for key in ("A", "B", "C", "D"):
                if key not in profiles:
                    profiles[key] = dict(base.get(key, {}))
        else:
            profiles = {k: dict(v) for k, v in base.items()}
        merge_scoring_into_profiles(profiles, scoring)
        data = {"questions": questions, "profiles": profiles}
        if scoring.get("intro"):
            data["scoring_intro"] = scoring["intro"]
        if scoring.get("mixed_results"):
            data["mixed_results"] = scoring["mixed_results"]
        out_path = OUT_DIR / f"{slug}.json"
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print("Wrote", out_path, "with", len(questions), "questions")

    # Generate understanding_data.json for all four assessments (+ engagement HTML include)
    generate_all_understanding_data()


if __name__ == "__main__":
    main()
