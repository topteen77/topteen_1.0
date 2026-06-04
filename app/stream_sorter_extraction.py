"""Extract post-table career guidance from Class 10 RIASEC Stream Sorter .docx files."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

STREAM_WISE_TITLE = 'Stream-Wise Premium Career Options'
FUTURE_TITLE = 'Most Future-Relevant Careers Across All Streams'

DOCX_FILES = [
    ('Realistic.docx', 'R'),
    ('Investigative.docx', 'I'),
    ('Artistic.docx', 'A'),
    ('social.docx', 'S'),
    ('Enterprising.docx', 'E'),
    ('Conventional.docx', 'C'),
]

STREAM_HEADER_RE = re.compile(
    r'^(PCM|PCB|CWM|CWOM|HUM|Commerce|Humanities|Fine Arts|Science)\b',
    re.IGNORECASE,
)


def _is_stream_header(text: str) -> bool:
    line = (text or '').strip()
    if not line or line == STREAM_WISE_TITLE:
        return False
    if STREAM_HEADER_RE.match(line):
        return True
    if '(' in line and re.search(
        r'\b(PCM|PCB|Physics|Chemistry|Mathematics|Biology|Commerce|Humanities|Math)\b',
        line,
        re.I,
    ):
        return True
    return False


def _paragraphs_after_table(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _parse_stream_wise_section(paras: list[str], start_idx: int, end_idx: int) -> list[dict]:
    groups: list[dict] = []
    current_stream = None
    current_careers: list[str] = []

    for text in paras[start_idx:end_idx]:
        if text == STREAM_WISE_TITLE:
            continue
        if _is_stream_header(text):
            if current_stream and current_careers:
                groups.append({'stream': current_stream, 'careers': current_careers})
            current_stream = text
            current_careers = []
            continue
        if current_stream:
            current_careers.append(text)

    if current_stream and current_careers:
        groups.append({'stream': current_stream, 'careers': current_careers})
    return groups


def extract_guidance_from_docx(docx_path: Path) -> dict:
    doc = Document(str(docx_path))
    paras = _paragraphs_after_table(doc)

    try:
        stream_start = next(i for i, p in enumerate(paras) if STREAM_WISE_TITLE in p)
    except StopIteration:
        stream_start = None
    try:
        future_start = next(i for i, p in enumerate(paras) if FUTURE_TITLE in p)
    except StopIteration:
        future_start = None

    stream_wise = []
    future_careers = []
    heading = paras[0] if paras else ''

    if stream_start is not None and future_start is not None and stream_start < future_start:
        stream_wise = _parse_stream_wise_section(paras, stream_start + 1, future_start)
        future_careers = [
            p for p in paras[future_start + 1:]
            if p and p != FUTURE_TITLE
        ]
    elif future_start is not None:
        future_careers = [
            p for p in paras[future_start + 1:]
            if p and p != FUTURE_TITLE
        ]

    category_codes = []
    if doc.tables:
        for row in doc.tables[0].rows[1:]:
            code = row.cells[0].text.strip().upper()
            if code:
                category_codes.append(code)

    return {
        'heading': heading,
        'category_codes': category_codes,
        'stream_wise_premium_careers': stream_wise,
        'future_relevant_careers': future_careers,
    }


def build_guidance_payload(source_dir: Path) -> dict:
    source_dir = Path(source_dir)
    files_data = {}
    category_code_to_letter: dict[str, str] = {}

    for filename, letter in DOCX_FILES:
        docx_path = source_dir / filename
        if not docx_path.exists():
            raise FileNotFoundError(f'Missing source file: {docx_path}')
        extracted = extract_guidance_from_docx(docx_path)
        files_data[filename] = {
            'riasec_letter': letter,
            **extracted,
        }
        for code in extracted['category_codes']:
            category_code_to_letter[code] = letter

    return {
        'version': 1,
        'source_directory': str(source_dir),
        'section_titles': {
            'stream_wise': STREAM_WISE_TITLE,
            'future_relevant': FUTURE_TITLE,
        },
        'files': files_data,
        'category_code_to_letter': category_code_to_letter,
    }
